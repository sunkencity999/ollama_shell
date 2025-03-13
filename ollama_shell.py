#!/usr/bin/env python3
import os
# Set environment variable to avoid tokenizers parallelism warning
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import json
import typer
import requests
import sys
import base64
from io import BytesIO
from PIL import Image
import mimetypes
import datetime
import markdown2
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration
from rich.console import Console, Group
from rich.markdown import Markdown
from rich.panel import Panel
from rich.prompt import Prompt
from rich.syntax import Syntax
from rich.table import Table
from rich.box import ASCII as ASCII_BOX
from rich.align import Align
from rich.live import Live
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TimeElapsedColumn
from pyfiglet import Figlet
from termcolor import colored
from typing import Optional, Union
from pathlib import Path
from duckduckgo_search import DDGS
from bs4 import BeautifulSoup
import concurrent.futures
import html2text
from prompt_toolkit import PromptSession
from prompt_toolkit.key_binding import KeyBindings
import pyperclip
from bs4 import BeautifulSoup
import time
import re

# Import vector database and embedding libraries
try:
    import chromadb
    from sentence_transformers import SentenceTransformer
    VECTOR_DB_AVAILABLE = True
except ImportError:
    VECTOR_DB_AVAILABLE = False

# Import file creation functions
from file_creation import create_file, create_text_file, create_csv_file, create_docx_file, create_excel_file, create_pdf_file

# Import fine-tuning module
try:
    # Try to import the new modular fine-tuning system
    from finetune_modules import FineTuningManager, detect_hardware
    FINETUNE_AVAILABLE = True
except ImportError:
    try:
        # Fall back to the old fine-tuning system for backward compatibility
        import finetune
        from finetune import FineTuningManager
        FINETUNE_AVAILABLE = True
    except ImportError:
        FINETUNE_AVAILABLE = False

# Import filesystem MCP integration
try:
    from filesystem_integration import FilesystemIntegration, handle_fs_command, filesystem_mode
    FILESYSTEM_AVAILABLE = True
except ImportError:
    FILESYSTEM_AVAILABLE = False

# Import filesystem MCP Protocol integration (for natural language commands)
try:
    from ollama_shell_filesystem_mcp import get_ollama_shell_filesystem_mcp, handle_filesystem_nl_command
    FILESYSTEM_MCP_AVAILABLE = True
except ImportError:
    FILESYSTEM_MCP_AVAILABLE = False

# Import Confluence MCP integration
try:
    from ollama_shell_confluence_mcp import get_ollama_shell_confluence_mcp, handle_confluence_nl_command, check_confluence_configuration, save_confluence_config, display_confluence_result
    CONFLUENCE_MCP_AVAILABLE = True
except ImportError:
    CONFLUENCE_MCP_AVAILABLE = False

# Import Jira MCP integration
try:
    from ollama_shell_jira_mcp import get_ollama_shell_jira_mcp, handle_jira_nl_command, check_jira_configuration, save_jira_config, display_jira_result
    JIRA_MCP_AVAILABLE = True
except ImportError:
    JIRA_MCP_AVAILABLE = False

app = typer.Typer()
console = Console(
    force_terminal=True,
    color_system="auto",
    highlight=True,
    legacy_windows=True  # Better handling of ANSI codes on Windows
)

OLLAMA_API = "http://localhost:11434/api"
CONFIG_FILE = os.path.join(os.path.dirname(__file__), "config.json")
DEFAULT_CONFIG = {
    "default_model": "llama3.2:latest",
    "default_vision_model": "llama3.2-vision:latest",
    "verbose": False,
    "save_history": True,
    "history_file": "~/.ollama_shell_history",
    "show_model_details": True,
    "temperature": 0.7,
    "context_length": 8192,
    "knowledge_base": {
        "enabled": True,
        "path": "~/.ollama_shell_kb",
        "max_results": 5,
        "similarity_threshold": 0.05
    },
    "stored_prompts": {
        "code_expert": {
            "title": "Code Review & Improvement Expert",
            "prompt": """You are an expert software developer with deep knowledge across many programming languages and best practices. Your task is to:
1. Review code for potential issues
2. Suggest improvements for readability and performance
3. Identify security concerns
4. Recommend modern alternatives to deprecated practices
5. Explain complex code sections
Be specific in your recommendations and include code examples when relevant."""
        },
        "technical_writer": {
            "title": "Technical Documentation Expert",
            "prompt": """You are a technical writing expert. Your task is to:
1. Create clear, concise documentation
2. Explain complex technical concepts in simple terms
3. Structure information logically
4. Include relevant examples and use cases
5. Highlight important warnings or prerequisites
Focus on clarity and completeness while maintaining a professional tone."""
        },
        "code_teacher": {
            "title": "Programming Teacher & Mentor",
            "prompt": """You are a patient and knowledgeable programming teacher. Your task is to:
1. Explain concepts from fundamental principles
2. Provide helpful analogies and examples
3. Break down complex topics into manageable pieces
4. Encourage best practices and good habits
5. Answer questions thoroughly and clearly
Remember to check understanding and provide additional examples when needed."""
        },
        "system_admin": {
            "title": "System Administration Expert",
            "prompt": """You are an experienced system administrator. Your task is to:
1. Provide clear command-line solutions
2. Explain system configurations and best practices
3. Troubleshoot system issues
4. Recommend security measures
5. Optimize system performance
Always consider security implications and provide warnings about potentially dangerous operations."""
        },
        "api_designer": {
            "title": "API Design Consultant",
            "prompt": """You are an API design expert. Your task is to:
1. Design clear and intuitive API endpoints
2. Follow REST/GraphQL best practices
3. Suggest appropriate data structures
4. Consider security and performance
5. Provide example requests and responses
Focus on creating APIs that are both powerful and developer-friendly."""
        }
    }
}

PROMPT_STYLE = None

# File creation helper functions
# Note: File creation functions are imported from file_creation.py

def load_config():
    """Load configuration from config file"""
    try:
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r') as f:
                return {**DEFAULT_CONFIG, **json.load(f)}
        return DEFAULT_CONFIG
    except Exception as e:
        console.print(f"[yellow]Warning: Could not load config, using defaults: {str(e)}[/yellow]")
        return DEFAULT_CONFIG

def save_config(config):
    """Save configuration to config file"""
    try:
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f, indent=4)
        console.print("[green]Configuration saved successfully[/green]")
    except Exception as e:
        console.print(f"[red]Error saving config: {str(e)}[/red]")

def save_history(model: str, messages: list):
    """Save chat history to file"""
    config = load_config()
    if not config["save_history"]:
        return
    
    history_file = os.path.expanduser(config["history_file"])
    try:
        history = []
        if os.path.exists(history_file):
            with open(history_file, 'r') as f:
                history = json.load(f)
        
        history.append({
            "timestamp": datetime.datetime.now().isoformat(),
            "model": model,
            "messages": messages
        })
        
        with open(history_file, 'w') as f:
            json.dump(history, f, indent=4)
    except Exception as e:
        console.print(f"[yellow]Warning: Could not save history: {str(e)}[/yellow]")

def display_banner():
    """Display the application banner"""
    try:
        f = Figlet(font='slant')
        banner = f.renderText('Ollama Shell')
    except Exception:
        # Fallback to simple banner if figlet fails
        banner = r"""
 ____  ____                         _____ __         ____ 
/ __ \/ / /___ _____ ___  ____ _  / ___// /_  ___  / / /
 / / / / / / __ `/ __ `__ \/ __ `/  \__ \/ __ \/ _ \/ / / 
/ /_/ / / / /_/ / / / / / / /_/ /  ___/ / / / /  __/ / /  
\____/_/_/\__,_/_/ /_/ /_/\__,_/  /____/_/ /_/\___/_/_/   
                                                           
"""
    
    # Clear the screen first
    console.clear()
    
    # Create panels with simple ASCII borders
    console.print(Panel(banner.rstrip(), border_style="cyan", padding=(0, 2)))
    console.print(Panel(":rocket: Your friendly command-line LLM interface", border_style="cyan", padding=(0, 2)))

def send_message(model: str, message: str, system_prompt: Optional[str] = None, stream: bool = True, context_messages: Optional[list] = None):
    """Enhanced message sending with configuration options and context support"""
    config = load_config()
    url = f"{OLLAMA_API}/chat"
    
    # If context_messages is provided, use it instead of creating a new message list
    if context_messages:
        payload = {
            "model": model,
            "messages": context_messages,
            "options": {
                "temperature": config["temperature"],
                "num_ctx": config["context_length"]
            },
            "stream": stream
        }
        
        # Add the user message if it's not already included
        if message and (not context_messages or context_messages[-1]["role"] != "user" or context_messages[-1]["content"] != message):
            payload["messages"].append({"role": "user", "content": message})
    else:
        # Traditional message format
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt} if system_prompt else None,
                {"role": "user", "content": message}
            ],
            "options": {
                "temperature": config["temperature"],
                "num_ctx": config["context_length"]
            },
            "stream": stream
        }
        # Remove None values from messages
        payload["messages"] = [msg for msg in payload["messages"] if msg is not None]
    
    try:
        response = requests.post(url, json=payload, stream=stream)
        response.raise_for_status()
        
        if stream:
            # Return the response object for streaming
            return response
        else:
            # Process non-streaming response
            try:
                result = response.json()
                if isinstance(result, dict) and "message" in result and "content" in result["message"]:
                    return result
                else:
                    raise Exception("Unexpected response format")
            except json.JSONDecodeError as e:
                # If JSON parsing fails, try to extract content from raw response
                content = response.text.strip()
                if content:
                    return {"message": {"content": content}}
                else:
                    raise Exception(f"Failed to parse response: {e}")
    except requests.exceptions.RequestException as e:
        raise Exception(f"Request failed: {e}")

def generate_chat_completion(model: str, messages: list) -> str:
    """Generate a chat completion using the specified model and messages"""
    config = load_config()
    url = f"{OLLAMA_API}/chat"
    
    payload = {
        "model": model,
        "messages": messages,
        "options": {
            "temperature": config["temperature"],
            "num_ctx": config["context_length"]
        },
        "stream": False
    }
    
    try:
        response = requests.post(url, json=payload)
        response.raise_for_status()
        
        # Process response
        try:
            result = response.json()
            if isinstance(result, dict) and "message" in result and "content" in result["message"]:
                return result["message"]["content"]
            else:
                raise Exception("Unexpected response format")
        except json.JSONDecodeError as e:
            # If JSON parsing fails, try to extract content from raw response
            content = response.text.strip()
            if content:
                return content
            else:
                raise Exception(f"Failed to parse response: {e}")
    except requests.exceptions.RequestException as e:
        raise Exception(f"Request failed: {e}")

def get_available_models():
    """Get list of available models from Ollama"""
    try:
        response = requests.get(f"{OLLAMA_API}/tags")
        if response.status_code == 200:
            return [model["name"] for model in response.json()["models"]]
        return []
    except Exception:
        return []

def validate_model(model: str) -> tuple[bool, str]:
    """Validate if a model exists and return (is_valid, error_message)"""
    try:
        available_models = get_available_models()
        if not available_models:
            return False, "Could not fetch available models. Is Ollama running?"
        
        if model not in available_models:
            message = f"Model '{model}' not found. Available models:\n"
            for i, m in enumerate(available_models, 1):
                message += f"  {i}. {m}\n"
            return False, message
        
        return True, ""
    except Exception as e:
        return False, f"Error validating model: {str(e)}"

def load_history():
    """Load chat history from file"""
    config = load_config()
    history_file = os.path.expanduser(config["history_file"])
    try:
        if os.path.exists(history_file):
            with open(history_file, 'r') as f:
                return json.load(f)
        return []
    except Exception as e:
        console.print(f"[yellow]Warning: Could not load history: {str(e)}[/yellow]")
        return []

def view_history():
    """View and manage chat history"""
    history = load_history()
    
    while True:
        console.clear()
        display_banner()
        
        if not history:
            console.print("[yellow]No chat history found[/yellow]")
            Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")
            return
        
        # Group chats by date
        from collections import defaultdict
        from datetime import datetime
        
        chats_by_date = defaultdict(list)
        for i, chat in enumerate(history):
            date = datetime.fromisoformat(chat["timestamp"]).strftime("%Y-%m-%d")
            chats_by_date[date].append((i, chat))
        
        # Display chats grouped by date
        table = Table(title="Chat History")
        table.add_column("Number", style="green")
        table.add_column("Date", style="cyan")
        table.add_column("Model", style="yellow")
        table.add_column("Messages", style="white")
        
        for date in sorted(chats_by_date.keys(), reverse=True):
            for i, chat in chats_by_date[date]:
                msg_count = len([m for m in chat["messages"] if m["role"] == "user"])
                preview = chat["messages"][1]["content"][:50] + "..." if len(chat["messages"]) > 1 else "Empty chat"
                table.add_row(
                    f"[green]{i + 1}[/green]",
                    date,
                    chat["model"],
                    f"{msg_count} messages - {preview}"
                )
        
        console.print(table)
        console.print("\n[cyan]Options:[/cyan]")
        console.print("[green]number[/green]: View chat details")
        console.print("[green]r number[/green]: Resume chat")
        console.print("[green]c[/green]: Clear all history")
        console.print("[green]x[/green]: Exit to main menu")
        console.print("[green]e number[/green]: Export chat")
        
        choice = Prompt.ask("\n[yellow]Enter your choice[/yellow]")
        
        if choice.lower() == 'x':
            break
        elif choice.lower() == 'c':
            confirm = Prompt.ask("\n[red]Are you sure you want to clear all history?[/red]", choices=["y", "n", "yes", "no"], default="n").lower()
            if confirm in ['y', 'yes']:
                config = load_config()
                history_file = os.path.expanduser(config["history_file"])
                try:
                    with open(history_file, 'w') as f:
                        json.dump([], f)
                    console.print("[green]History cleared successfully[/green]")
                    history = []
                except Exception as e:
                    console.print(f"[red]Error clearing history: {str(e)}[/red]")
            Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")
        elif choice.startswith('r ') and choice[2:].isdigit():
            chat_num = int(choice[2:])
            if 1 <= chat_num <= len(history):
                chat = history[chat_num - 1]
                console.clear()
                display_banner()
                
                # Resume chat with the same model and history
                console.print(f"\n[cyan]Resuming chat from {chat['timestamp']}[/cyan]")
                console.print(f"[yellow]Model: {chat['model']}[/yellow]\n")
                
                # Display previous messages
                for msg in chat["messages"]:
                    if msg["role"] == "system":
                        console.print(f"[yellow]System: {msg['content']}[/yellow]")
                    elif msg["role"] == "user":
                        console.print(f"\n[cyan]You:[/cyan]")
                        console.print(Panel(msg["content"], border_style="cyan"))
                    else:
                        console.print(f"\n[purple]Assistant:[/purple]")
                        console.print(Panel(Markdown(msg["content"]), border_style="purple"))
                
                # Start interactive chat with existing history
                interactive_chat(chat["model"], 
                              system_prompt=next((msg["content"] for msg in chat["messages"] if msg["role"] == "system"), None),
                              existing_history=chat["messages"])
                break
        elif choice.startswith('e ') and choice[2:].isdigit():
            chat_num = int(choice[2:])
            if 1 <= chat_num <= len(history):
                chat = history[chat_num - 1]
                console.clear()
                display_banner()
                
                # Export chat
                console.print(f"\n[cyan]Exporting chat from {chat['timestamp']}[/cyan]")
                console.print(f"[yellow]Model: {chat['model']}[/yellow]\n")
                
                # Ask for export format
                console.print("\n[cyan]Select export format:[/cyan]")
                console.print("[green]1[/green]. Markdown")
                console.print("[green]2[/green]. HTML")
                console.print("[green]3[/green]. PDF")
                
                format_choice = Prompt.ask("\n[yellow]Enter your choice[/yellow]", choices=["1", "2", "3"], default="1")
                
                if format_choice == "1":
                    format = "markdown"
                elif format_choice == "2":
                    format = "html"
                elif format_choice == "3":
                    format = "pdf"
                else:
                    console.print("[red]Invalid choice. Defaulting to markdown.[/red]")
                    format = "markdown"
                
                # Ask for output file
                output_file = Prompt.ask("\n[cyan]Enter output file name (or press Enter for default)[/cyan]", default="")
                
                # Confirm export
                confirm = Prompt.ask("\n[yellow]Proceed with export?[/yellow]", choices=["y", "n", "yes", "no"], default="y").lower()
                if confirm in ['y', 'yes']:
                    try:
                        output_path = export_chat(chat["messages"], format, output_file)
                        console.print(f"\n[green]Chat exported successfully to {output_path}![/green]")
                    except Exception as e:
                        console.print(f"\n[red]Error exporting chat: {str(e)}[/red]")
                else:
                    console.print("\n[yellow]Export cancelled.[/yellow]")
                
                Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")
        elif choice.isdigit() and 1 <= int(choice) <= len(history):
            chat = history[int(choice) - 1]
            console.clear()
            display_banner()
            
            # Display chat details
            console.print(f"\n[cyan]Chat from {chat['timestamp']}[/cyan]")
            console.print(f"[yellow]Model: {chat['model']}[/yellow]\n")
            
            for msg in chat["messages"]:
                if msg["role"] == "system":
                    console.print(f"[yellow]System: {msg['content']}[/yellow]")
                elif msg["role"] == "user":
                    console.print(f"\n[cyan]You:[/cyan]")
                    console.print(Panel(msg["content"], border_style="cyan"))
                else:
                    console.print(f"\n[purple]Assistant:[/purple]")
                    console.print(Panel(Markdown(msg["content"]), border_style="purple"))
            
            Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")

def export_chat(chat_history: list, format: str = "markdown", output_file: str = None) -> str:
    """Export chat history to various formats
    
    Args:
        chat_history: List of chat messages
        format: Output format (markdown, html, or pdf)
        output_file: Optional output file path
    
    Returns:
        str: Path to the exported file
    """
    if not chat_history:
        raise ValueError("No chat history to export")
    
    # Generate timestamp for default filename
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Convert chat history to markdown
    markdown_content = "# Chat History\n\n"
    markdown_content += f"Exported on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    
    for msg in chat_history:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        
        if role == "system":
            markdown_content += f"### System\n{content}\n\n"
        elif role == "user":
            markdown_content += f"### You\n{content}\n\n"
        elif role == "assistant":
            markdown_content += f"### Assistant\n{content}\n\n"
    
    # Create output directory if it doesn't exist
    output_dir = os.path.expanduser("~/ollama_shell_exports")
    os.makedirs(output_dir, exist_ok=True)
    
    # Determine output path
    if not output_file:
        output_file = os.path.join(output_dir, f"chat_export_{timestamp}.{format}")
    
    output_path = os.path.expanduser(output_file)
    
    try:
        if format == "markdown":
            # Save as markdown
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
        
        elif format == "html":
            # Convert markdown to HTML
            import markdown2
            html_content = markdown2.markdown(markdown_content, extras=["fenced-code-blocks", "tables"])
            
            # Add some basic styling
            html_template = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
                    h1 {{ color: #2c3e50; }}
                    h3 {{ color: #34495e; margin-top: 20px; }}
                    pre {{ background-color: #f7f9fa; padding: 10px; border-radius: 5px; }}
                    code {{ font-family: Monaco, monospace; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_template)
        
        elif format == "pdf":
            # First convert to HTML, then to PDF
            import markdown2
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
            
            html_content = markdown2.markdown(markdown_content, extras=["fenced-code-blocks", "tables"])
            
            # Add some basic styling with proper font configuration
            font_config = FontConfiguration()
            css = CSS(string='''
                @page {
                    size: letter;
                    margin: 2cm;
                }
                body {
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
                }
                h1 {
                    color: #2c3e50;
                    font-size: 24px;
                    margin-bottom: 20px;
                }
                h3 {
                    color: #34495e;
                    font-size: 18px;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }
                pre {
                    background-color: #f7f9fa;
                    padding: 10px;
                    border-radius: 5px;
                    font-family: "SF Mono", Consolas, "Liberation Mono", Menlo, Courier, monospace;
                    font-size: 14px;
                    white-space: pre-wrap;
                }
                code {
                    font-family: "SF Mono", Consolas, "Liberation Mono", Menlo, Courier, monospace;
                    font-size: 14px;
                }
            ''', font_config=font_config)
            
            html_template = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Create PDF with proper font configuration
            HTML(string=html_template).write_pdf(output_path, stylesheets=[css], font_config=font_config)
        
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        return output_path
            
    except Exception as e:
        raise Exception(f"Error exporting chat: {str(e)}")

def get_model_context_window(model: str) -> int:
    """Get the context window size for a specific model"""
    # First check if user has specified a context length in config
    config = load_config()
    user_context_length = config.get("context_length")
    
    # If user has specified a context length, use that
    if user_context_length:
        return user_context_length
        
    # Otherwise try to get from Ollama API
    try:
        response = requests.get(f"{OLLAMA_API}/show", params={"name": model})
        if response.status_code == 200:
            model_info = response.json()
            # Extract context window from model parameters
            if "parameters" in model_info:
                return model_info["parameters"].get("num_ctx", 4096)  # Default to 4096 if not specified
        return 4096  # Default fallback
    except Exception:
        return 4096  # Default fallback

def count_tokens(text: str, model: str = None) -> int:
    """Count tokens in a text string"""
    # Simple estimation - can be improved with tiktoken or similar
    return len(text.split()) * 1.3  # Rough approximation

def format_token_count(current_tokens: int, max_tokens: int) -> str:
    """Format token count for display"""
    percentage = (current_tokens / max_tokens) * 100
    return f"Tokens: {current_tokens}/{max_tokens} ({percentage:.1f}%)"

class KnowledgeBase:
    """Knowledge base using vector embeddings for semantic search"""
    
    def __init__(self, path: str = None):
        """Initialize the knowledge base"""
        if not VECTOR_DB_AVAILABLE:
            raise ImportError("Vector database dependencies not available. Install with: pip install chromadb sentence-transformers")
        
        self.config = load_config()
        kb_config = self.config.get("knowledge_base", {})
        
        # Use provided path or default from config
        self.path = path or os.path.expanduser(kb_config.get("path", "~/.ollama_shell_kb"))
        self.max_results = kb_config.get("max_results", 5)
        self.similarity_threshold = kb_config.get("similarity_threshold", 0.05)
        
        # Create directory if it doesn't exist
        os.makedirs(self.path, exist_ok=True)
        
        # Initialize the vector database
        self.client = chromadb.PersistentClient(path=self.path)
        
        # Create or get the collection
        try:
            self.collection = self.client.get_collection("ollama_shell_kb")
        except ValueError:
            self.collection = self.client.create_collection(
                "ollama_shell_kb",
                metadata={"description": "Ollama Shell Knowledge Base"}
            )
        
        # Initialize the embedding model
        self.embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    
    def add_to_knowledge_base(self, text: str, metadata: dict = None, id_prefix: str = None) -> str:
        """Add text to the knowledge base"""
        if not text or len(text.strip()) < 10:
            return None
            
        # Generate a unique ID
        doc_id = f"{id_prefix or 'doc'}_{int(time.time())}_{hash(text) % 10000}"
        
        # Default metadata if none provided
        if metadata is None:
            metadata = {"source": "chat", "timestamp": datetime.datetime.now().isoformat()}
        
        try:
            # Generate embedding using the model
            embedding = self.embedding_model.encode(text).tolist()
            
            # Add to collection
            self.collection.add(
                documents=[text],
                embeddings=[embedding],
                metadatas=[metadata],
                ids=[doc_id]
            )
            
            config = load_config()
            if config.get("verbose", False):
                print(f"Debug - Added document with ID: {doc_id}")
                print(f"Debug - Collection count after add: {self.collection.count()}")
            
            return doc_id
        except Exception as e:
            config = load_config()
            if config.get("verbose", False):
                print(f"Debug - Error in add_to_knowledge_base: {str(e)}")
            return None
    
    def search_knowledge_base(self, query: str, limit: int = None) -> list:
        """Search the knowledge base for relevant information"""
        if not query or len(query.strip()) < 3:
            return []
            
        # Generate embedding for the query
        query_embedding = self.embedding_model.encode(query).tolist()
        
        # Search the collection
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=limit or self.max_results
        )
        
        # Debug output only when verbose mode is enabled
        config = load_config()
        if config.get("verbose", False):
            print(f"Debug - Query: '{query}'")
            print(f"Debug - Collection count: {self.collection.count()}")
            print(f"Debug - Raw results: {results.keys()}")
            if "documents" in results:
                print(f"Debug - Documents found: {len(results['documents'][0]) if results['documents'] else 0}")
                if results["documents"] and len(results["documents"][0]) > 0:
                    print(f"Debug - Distances: {results['distances'][0] if 'distances' in results else 'N/A'}")
                    print(f"Debug - Similarity threshold: {self.similarity_threshold}")
        
        # Format results
        formatted_results = []
        if results and "documents" in results and results["documents"]:
            for i, doc in enumerate(results["documents"][0]):
                # Get metadata and distance if available
                metadata = results["metadatas"][0][i] if "metadatas" in results and results["metadatas"] else {}
                distance = results["distances"][0][i] if "distances" in results and results["distances"] else 1.0
                
                # Only include results above similarity threshold
                similarity = 1.0 - distance  # Convert distance to similarity
                
                # Debug output for each result only when verbose mode is enabled
                config = load_config()
                if config.get("verbose", False):
                    print(f"Debug - Result {i} similarity: {similarity:.4f}")
                
                if similarity >= self.similarity_threshold:
                    formatted_results.append({
                        "text": doc,
                        "metadata": metadata,
                        "similarity": similarity
                    })
        
        return formatted_results
    
    def delete_from_knowledge_base(self, doc_id: str) -> bool:
        """Delete a document from the knowledge base by ID"""
        try:
            self.collection.delete(ids=[doc_id])
            return True
        except Exception:
            return False
    
    def get_stats(self) -> dict:
        """Get statistics about the knowledge base"""
        count = self.collection.count()
        return {
            "count": count,
            "path": self.path
        }
        
    def reset_knowledge_base(self) -> bool:
        """Delete all documents from the knowledge base"""
        try:
            # Get all document IDs
            all_ids = self.collection.get()['ids']
            
            if all_ids:
                # Delete all documents
                self.collection.delete(ids=all_ids)
                
            return True
        except Exception as e:
            print(f"Error resetting knowledge base: {str(e)}")
            return False
        
    def add_document(self, content: str, source: str, file_type: str) -> dict:
        """Add a document to the knowledge base, chunking if necessary
        
        Args:
            content: The document content
            source: The document source (usually filename)
            file_type: The type of file
            
        Returns:
            dict: Statistics about the operation
        """
        if not content or len(content.strip()) < 50:
            return {"success": False, "reason": "Document too short"}
            
        # For large documents, split into chunks of ~1000 characters with overlap
        content_length = len(content)
        chunk_size = 1000
        overlap = 100
        chunks_added = 0
        
        # Debug output only when verbose mode is enabled
        config = load_config()
        if config.get("verbose", False):
            print(f"Debug - Adding document: {source}")
            print(f"Debug - Content length: {content_length}")
            print(f"Debug - Collection count before: {self.collection.count()}")
        
        try:
            if content_length <= chunk_size:
                # Small document, add as single entry
                doc_id = self.add_to_knowledge_base(
                    content,
                    metadata={
                        "source": source,
                        "type": file_type,
                        "timestamp": datetime.datetime.now().isoformat()
                    },
                    id_prefix=f"doc_{source.replace(' ', '_')}"
                )
                if doc_id:
                    chunks_added = 1
            else:
                # Large document, split into chunks
                chunks = []
                for i in range(0, content_length, chunk_size - overlap):
                    chunk = content[i:i + chunk_size]
                    if len(chunk) > 100:  # Only add substantial chunks
                        chunks.append(chunk)
                
                # Add each chunk to knowledge base
                for i, chunk in enumerate(chunks):
                    doc_id = self.add_to_knowledge_base(
                        chunk,
                        metadata={
                            "source": source,
                            "type": file_type,
                            "chunk": i + 1,
                            "total_chunks": len(chunks),
                            "timestamp": datetime.datetime.now().isoformat()
                        },
                        id_prefix=f"doc_{source.replace(' ', '_')}_chunk{i+1}"
                    )
                    if doc_id:
                        chunks_added += 1
            
            return {
                "success": chunks_added > 0,
                "chunks_added": chunks_added,
                "total_length": content_length
            }
        except Exception as e:
            config = load_config()
            if config.get("verbose", False):
                print(f"Debug - Error adding document: {str(e)}")
            return {"success": False, "reason": str(e)}

# Initialize the knowledge base if available
kb_instance = None
if VECTOR_DB_AVAILABLE:
    try:
        config = load_config()
        if config.get("verbose", False):
            print("Debug - Initializing knowledge base...")
        kb_instance = KnowledgeBase()
        if config.get("verbose", False):
            print(f"Debug - Knowledge base initialized with {kb_instance.collection.count()} documents")
    except Exception as e:
        console.print(f"[yellow]Warning: Could not initialize knowledge base: {str(e)}[/yellow]")
        config = load_config()
        if config.get("verbose", False):
            print(f"Debug - Knowledge base initialization error: {str(e)}")

def interactive_chat(model: str, system_prompt: Optional[str] = None, context_files: Optional[list[str]] = None, existing_history: Optional[list] = None):
    """Start an interactive chat session with the specified model"""
    config = load_config()
    max_tokens = get_model_context_window(model)
    current_tokens = 0
    
    # Create separator styles
    copy_text = "(Ctrl+C to copy)"
    separator_width = console.width - 2 - len(copy_text) - 1  # -2 for margins, -1 for space
    separator = "─" * separator_width
    
    user_separator = f"[blue]{separator}[/blue]"
    assistant_separator = f"[green]{separator}[/green] [cyan]{copy_text}[/cyan]"
    
    # Validate model before starting chat
    is_valid, error_message = validate_model(model)
    if not is_valid:
        console.print(f"\n[red]Error: {error_message}[/red]")
        return
    
    console.clear()
    display_banner()
    console.print(f"\n[green]Starting chat with model: [bold]{model}[/bold][/green]")
    
    # Prepare spinner for loading model
    with Progress(
        SpinnerColumn(spinner_name="dots"),
        TextColumn("[bold green]Loading model..."),
        BarColumn(bar_width=40, complete_style="green", finished_style="green"),
        TimeElapsedColumn(),
        transient=True,
    ) as progress:
        task = progress.add_task("load", total=100)
        
        # Simulate progress updates while waiting for model to load
        for i in range(0, 101, 5):
            progress.update(task, completed=i)
            time.sleep(0.05)
    
    # Display existing chat history if present
    if existing_history:
        console.print("\n[yellow]Resuming previous chat...[/yellow]")
        for message in existing_history:
            if message["role"] == "user":
                console.print("\nYou:", style="cyan")
                console.print(f"  {message['content']}")
            elif message["role"] == "assistant":
                console.print("\n[green]Assistant:[/green]")
                console.print(Markdown(message["content"]))
    
    # Set up key bindings for drag-and-drop and clipboard
    kb = KeyBindings()
    drag_drop_active = False

    class ChatState:
        def __init__(self):
            self.document_context = ""
            self.pinned_messages = []  # Store pinned messages
            self.excluded_messages = []  # Store indices of messages to exclude
            self.summary = None  # Store conversation summary
            self.kb_enabled = VECTOR_DB_AVAILABLE and load_config().get("knowledge_base", {}).get("enabled", True)
            
    chat_state = ChatState()

    @kb.add('c-v')
    def _(event):
        nonlocal drag_drop_active
        drag_drop_active = not drag_drop_active
        if drag_drop_active:
            console.print("\n[cyan]Drag & Drop mode activated (Ctrl+V to toggle)[/cyan]")
            console.print("[cyan]Drag a file into the terminal...[/cyan]")
        else:
            console.print("\n[cyan]Drag & Drop mode deactivated[/cyan]")
    
    @kb.add('c-c')
    def _(event):
        nonlocal last_response
        if last_response:
            pyperclip.copy(last_response)
            console.print("\n[cyan]Last response copied to clipboard![/cyan]")
        else:
            console.print("\n[yellow]No response to copy[/yellow]")
    
    # Create prompt session with key bindings
    session = PromptSession(key_bindings=kb)
    
    # Initialize base system prompt
    base_system_prompt = """You are a helpful assistant that can engage in general conversation while also referencing 
    documents that have been shared. When discussing shared documents, be specific and reference their content when relevant.
    However, you can also discuss other topics naturally."""
    
    if system_prompt:
        base_system_prompt = system_prompt + "\n\n" + base_system_prompt
    
    # Prepare initial document context if files are provided
    if context_files:
        try:
            chat_state.document_context = prepare_document_context(context_files)
            console.print(f"[yellow]Loaded {len(context_files)} document(s) into context[/yellow]")
        except Exception as e:
            console.print(f"[yellow]Warning: Could not read {context_files}: {str(e)}[/yellow]")
            return

    # Function to get current system prompt with document context
    def get_current_system_prompt():
        current_prompt = base_system_prompt
        if chat_state.document_context:
            current_prompt += f"\n\nAvailable document context:\n{chat_state.document_context}"
        return current_prompt

    if not system_prompt and config["stored_prompts"]:
        # Offer to use a stored prompt
        console.print("\n[bold yellow]Stored System Prompts[/bold yellow]")
        prompts = list(config["stored_prompts"].items())
        
        # Add "none" option
        console.print("[green]0[/green]. No system prompt")
        
        # List available prompts
        for i, (name, prompt) in enumerate(prompts, 1):
            preview = prompt["prompt"][:50] + "..." if len(prompt["prompt"]) > 50 else prompt["prompt"]
            console.print(f"[green]{i}[/green]. {name}: {preview}")
        
        choice = Prompt.ask("\nSelect a system prompt", choices=[str(i) for i in range(len(prompts) + 1)])
        choice = int(choice)
        
        if choice > 0:
            base_system_prompt = prompts[choice - 1][1]["prompt"] + "\n\n" + base_system_prompt
            console.print(f"\n[cyan]Using stored prompt: {prompts[choice - 1][0]}[/cyan]")
    
    chat_history = existing_history if existing_history else []
    chat_history.append({"role": "system", "content": get_current_system_prompt()})

    # Display chat instructions in a styled panel
    instructions = "Type 'exit' to end.\n• Press Ctrl+V to toggle drag & drop mode for file sharing\n• Press Ctrl+C to copy assistant responses\n• Type '/help' to see context management commands"
    console.print(Panel(instructions, title="[bold cyan]Chat Session Started[/bold cyan]", border_style="cyan", padding=(1, 2)))

    last_response = ""

    while True:
        try:
            # Use prompt_toolkit session for input
            console.print(user_separator)
            console.print("\nYou:", style="cyan")
            user_input = session.prompt("  ").strip()  # Two spaces for proper alignment
            
            if user_input.lower() in ['exit', 'quit', 'q']:
                # Offer to export chat before exiting
                if chat_history and Prompt.ask("\nWould you like to export this chat? [y/N]", choices=["y", "n", "yes", "no"], default="n").lower() in ['y', 'yes']:
                    format_choice = Prompt.ask(
                        "Choose format",
                        choices=["markdown", "html", "pdf"],
                        default="markdown"
                    )
                    try:
                        output_file = export_chat(chat_history, format=format_choice)
                        console.print(f"[green]Chat exported to: {output_file}![/green]")
                    except Exception as e:
                        console.print(f"[red]Error exporting chat: {str(e)}[/red]")
                break
                
            # Handle drag & drop mode first
            if drag_drop_active and user_input:
                # Clean up the file path by removing escape characters and expanding user path
                cleaned_path = os.path.expanduser(user_input.replace('\\', ''))
                if os.path.exists(cleaned_path):
                    try:
                        content, file_type = read_file_content(cleaned_path)
                        # Update document context
                        chat_state.document_context += f"\n\nContent from {os.path.basename(cleaned_path)} ({file_type}):\n{content}"
                        
                        # Check if we should add to knowledge base
                        add_to_kb = False
                        if VECTOR_DB_AVAILABLE and kb_instance and chat_state.kb_enabled:
                            # Skip knowledge base prompt for image files
                            _, ext = os.path.splitext(cleaned_path)
                            ext = ext.lower()
                            is_image = ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
                            
                            if not is_image:
                                add_to_kb = Prompt.ask(
                                    f"\nAdd [cyan]{os.path.basename(cleaned_path)}[/cyan] to knowledge base?",
                                    choices=["y", "n", "yes", "no"],
                                    default="n"
                                ).lower() in ['y', 'yes']
                        
                        # Create a message about the file being shared
                        file_message = f"I've added a {file_type} to our conversation. Please analyze it and provide key insights. The content is:\n\n{content}"
                        
                        # Send message to model (non-streaming for file analysis)
                        with console.status("[cyan]Analyzing document...[/cyan]"):
                            try:
                                response = send_message(model, file_message, get_current_system_prompt(), stream=False)
                                if response and "message" in response and "content" in response["message"]:
                                    assistant_message = response["message"]["content"]
                                    chat_history.append({"role": "user", "content": file_message})
                                    chat_history.append({"role": "assistant", "content": assistant_message})
                                    console.print("\n[green]Successfully loaded file: {os.path.basename(cleaned_path)}[/green]")
                                    
                                    # Add to knowledge base if requested
                                    if add_to_kb:
                                        console.print("[cyan]Adding to knowledge base...[/cyan]")
                                        try:
                                            result = kb_instance.add_document(
                                                content=content,
                                                source=os.path.basename(cleaned_path),
                                                file_type=file_type
                                            )
                                            
                                            if result["success"]:
                                                console.print(f"[green]Added document to knowledge base in {result['chunks_added']} chunks[/green]")
                                            else:
                                                console.print(f"[yellow]Failed to add to knowledge base: {result.get('reason', 'Unknown error')}[/yellow]")
                                        except Exception as e:
                                                   console.print(f"[red]Error adding to knowledge base: {str(e)}[/red]")
                                    console.print("\n[green]Assistant:[/green]")
                                    console.print(Panel(Markdown(assistant_message), border_style="purple"))
                                else:
                                    console.print("[red]Error: Unexpected response format from model[/red]")
                            except Exception as e:
                                console.print(f"[red]Error getting model response: {str(e)}[/red]")
                    except Exception as e:
                        console.print(f"[red]Error reading file: {str(e)}[/red]")
                    continue
                
            # Handle context management commands
            if user_input.startswith('/'):
                command_parts = user_input.split(' ', 1)
                command = command_parts[0].lower()
                args = command_parts[1] if len(command_parts) > 1 else ""
                
                if command == '/help':
                    console.print("\n[bold cyan]Context Management Commands:[/bold cyan]")
                    console.print("[green]/pin [message_number][/green] - Pin a message to keep it in context")
                    console.print("[green]/unpin [message_number][/green] - Unpin a previously pinned message")
                    console.print("[green]/exclude [message_number][/green] - Exclude a message from context")
                    console.print("[green]/include [message_number][/green] - Include a previously excluded message")
                    console.print("[green]/summarize[/green] - Summarize the conversation to save tokens")
                    console.print("[green]/context[/green] - Show current context management status")
                    console.print("[green]/tokens[/green] - Show token usage information")
                    console.print("[green]/help[/green] - Show this help message")
                    
                    console.print("\n[bold cyan]Knowledge Base Commands:[/bold cyan]")
                    console.print("[green]/kb status[/green] - Show knowledge base status")
                    console.print("[green]/kb add [text][/green] - Add text to knowledge base")
                    console.print("[green]/kb search [query][/green] - Search knowledge base")
                    console.print("[green]/kb toggle[/green] - Enable/disable knowledge base")
                    
                    if FINETUNE_AVAILABLE:
                        console.print("\n[bold cyan]Fine-Tuning Commands:[/bold cyan]")
                        console.print("[green]/finetune status[/green] - Show fine-tuning status and hardware detection")
                        console.print("[green]/finetune create [name] [base_model][/green] - Create a new fine-tuning job")
                        console.print("[green]/finetune dataset [path][/green] - Prepare a dataset for fine-tuning")
                        console.print("[green]/finetune start [name][/green] - Start a fine-tuning job")
                        console.print("[green]/finetune list[/green] - List all fine-tuning jobs")
                        console.print("[green]/finetune export [name][/green] - Export model to Ollama")
                        console.print("[green]/finetune install[/green] - Install fine-tuning dependencies")
                        console.print("[green]/finetune pause [name][/green] - Pause a running fine-tuning job")
                        console.print("[green]/finetune resume [name][/green] - Resume a paused fine-tuning job")
                        console.print("[green]/finetune delete [name][/green] - Delete a fine-tuning job")
                        console.print("[green]/finetune models[/green] - List available Ollama models for fine-tuning")
                        console.print("[green]/finetune progress [name][/green] - Show progress of a fine-tuning job")
                        console.print("[green]/finetune datasets[/green] - List available datasets")
                        console.print("[green]/finetune dataset-set [job_name] [dataset_id][/green] - Update dataset for a job")
                        console.print("[green]/finetune dataset-remove [dataset_id] [--force][/green] - Remove a dataset")
                        console.print("[green]/finetune reset [name][/green] - Reset a job to created state")
                    
                    if CONFLUENCE_MCP_AVAILABLE:
                        console.print("\n[bold cyan]Confluence Commands:[/bold cyan]")
                        console.print("[green]/confluence [command][/green] - Execute Confluence operations using natural language")
                        console.print("[green]/confluence mode[/green] - Enter interactive Confluence mode")
                    
                    console.print("\n[bold cyan]Other Commands:[/bold cyan]")
                    console.print("[green]search: [query][/green] - Perform a web search and analyze results")
                    console.print("[green]/fsnl [command][/green] - Execute filesystem operations using natural language")
                    console.print("[green]/create [filename] [content][/green] - Create a file with specified content")
                    console.print("[green]/create [request] and save to [filename][/green] - Generate and save content")
                    console.print("[green]Ctrl+V[/green] - Toggle drag & drop mode for file sharing")
                    console.print("[green]Ctrl+C[/green] - Copy assistant's response to clipboard")
                    console.print("[green]exit, quit, q[/green] - Exit the chat")
                    console.print("[green]e number[/green]: Export chat")
                    
                    continue
                    
                elif command == '/pin':
                    if args and args.strip().isdigit():
                        msg_idx = int(args.strip())
                        if msg_idx in chat_state.pinned_messages:
                            console.print(f"[yellow]Message {msg_idx} is already pinned[/yellow]")
                        elif 0 < msg_idx < len(chat_history):
                            chat_state.pinned_messages.append(msg_idx)
                            console.print(f"[green]Message {msg_idx} pinned to context[/green]")
                        else:
                            console.print(f"[red]Invalid message number. Must be between 1 and {len(chat_history)-1}[/red]")
                    else:
                        console.print("[red]Please specify a valid message number to pin (e.g., /pin 3)[/red]")
                    continue
                    
                elif command == '/unpin':
                    if args and args.strip().isdigit():
                        msg_idx = int(args.strip())
                        if msg_idx in chat_state.pinned_messages:
                            chat_state.pinned_messages.remove(msg_idx)
                            console.print(f"[green]Message {msg_idx} unpinned[/green]")
                        else:
                            console.print(f"[yellow]Message {msg_idx} is not pinned[/yellow]")
                    else:
                        console.print("[red]Please specify a valid message number to unpin (e.g., /unpin 3)[/red]")
                    continue
                    
                elif command == '/exclude':
                    if args and args.strip().isdigit():
                        msg_idx = int(args.strip())
                        if 0 < msg_idx < len(chat_history):
                            if msg_idx not in chat_state.excluded_messages:
                                chat_state.excluded_messages.append(msg_idx)
                                console.print(f"[green]Message {msg_idx} excluded from context[/green]")
                            else:
                                console.print(f"[yellow]Message {msg_idx} is already excluded[/yellow]")
                        else:
                            console.print(f"[red]Invalid message number. Must be between 1 and {len(chat_history)-1}[/red]")
                    else:
                        console.print("[red]Please specify a valid message number to exclude (e.g., /exclude 3)[/red]")
                    continue
                    
                elif command == '/include':
                    if args and args.strip().isdigit():
                        msg_idx = int(args.strip())
                        if msg_idx in chat_state.excluded_messages:
                            chat_state.excluded_messages.remove(msg_idx)
                            console.print(f"[green]Message {msg_idx} included in context[/green]")
                        else:
                            console.print(f"[yellow]Message {msg_idx} is not excluded[/yellow]")
                    else:
                        console.print("[red]Please specify a valid message number to include (e.g., /include 3)[/red]")
                    continue
                    
                elif command == '/summarize':
                    # Generate a summary of the conversation
                    with console.status("[cyan]Generating conversation summary...[/cyan]"):
                        summary_prompt = """Please create a concise summary of the conversation so far. 
                        Focus on the key points, questions, and answers. This summary will be used to 
                        maintain context while saving tokens."""
                        
                        # Prepare a temporary message list for the summary request
                        summary_messages = [
                            {"role": "system", "content": "You are a helpful assistant that summarizes conversations."},
                            {"role": "user", "content": summary_prompt}
                        ]
                        
                        # Add the conversation history excluding system messages
                        for msg in chat_history:
                            if msg["role"] != "system":
                                summary_messages.append(msg)
                        
                        try:
                            response = send_message(model, summary_prompt, "You are a helpful assistant that summarizes conversations.", stream=False)
                            if response and "message" in response and "content" in response["message"]:
                                chat_state.summary = response["message"]["content"]
                                console.print("\n[green]Conversation summarized successfully[/green]")
                                console.print(Panel(Markdown(chat_state.summary), title="Conversation Summary", border_style="green"))
                            else:
                                console.print("[red]Error: Failed to generate summary[/red]")
                        except Exception as e:
                            console.print(f"[red]Error generating summary: {str(e)}[/red]")
                    continue
                    
                elif command == '/context':
                    # Display context management status
                    console.print("\n[bold cyan]Context Management Status:[/bold cyan]")
                    
                    # Count tokens
                    total_tokens = sum(count_tokens(msg["content"], model) for msg in chat_history)
                    max_tokens = config["context_length"]
                    
                    # Display pinned messages
                    if chat_state.pinned_messages:
                        console.print("\n[yellow]Pinned Messages:[/yellow]")
                        for idx in chat_state.pinned_messages:
                            if 0 <= idx < len(chat_history):
                                msg = chat_history[idx]
                                preview = msg["content"][:50] + "..." if len(msg["content"]) > 50 else msg["content"]
                                console.print(f"[green]{idx}[/green]: [{msg['role']}] {preview}")
                    else:
                        console.print("\n[yellow]No pinned messages[/yellow]")
                    
                    # Display excluded messages
                    if chat_state.excluded_messages:
                        console.print("\n[yellow]Excluded Messages:[/yellow]")
                        for idx in chat_state.excluded_messages:
                            if 0 <= idx < len(chat_history):
                                msg = chat_history[idx]
                                preview = msg["content"][:50] + "..." if len(msg["content"]) > 50 else msg["content"]
                                console.print(f"[green]{idx}[/green]: [{msg['role']}] {preview}")
                    else:
                        console.print("\n[yellow]No excluded messages[/yellow]")
                    
                    # Display summary status
                    if chat_state.summary:
                        console.print("\n[yellow]Conversation Summary:[/yellow]")
                        console.print(Panel(Markdown(chat_state.summary[:200] + "..." if len(chat_state.summary) > 200 else chat_state.summary), 
                                           border_style="green"))
                    else:
                        console.print("\n[yellow]No conversation summary[/yellow]")
                    
                    # Display token usage
                    console.print(f"\n[yellow]Total tokens: {total_tokens}/{max_tokens} ({(total_tokens/max_tokens)*100:.1f}%)[/yellow]")
                    continue
                    
                elif command == '/tokens':
                    # Display detailed token usage
                    console.print("\n[bold cyan]Token Usage:[/bold cyan]")
                    max_tokens = config["context_length"]
                    
                    table = Table(title="Message Token Counts")
                    table.add_column("Msg #", style="cyan")
                    table.add_column("Role", style="green")
                    table.add_column("Tokens", style="yellow")
                    table.add_column("Status", style="magenta")
                    table.add_column("Preview", style="white")
                    
                    total_tokens = 0
                    for i, msg in enumerate(chat_history):
                        tokens = count_tokens(msg["content"], model)
                        total_tokens += tokens
                        
                        status = []
                        if i in chat_state.pinned_messages:
                            status.append("Pinned")
                        if i in chat_state.excluded_messages:
                            status.append("Excluded")
                        status_str = ", ".join(status) if status else "Normal"
                        
                        preview = msg["content"][:30] + "..." if len(msg["content"]) > 30 else msg["content"]
                        preview = preview.replace("\n", " ")
                        
                        table.add_row(
                            str(i),
                            msg["role"],
                            str(tokens),
                            status_str,
                            preview
                        )
                    
                    console.print(table)
                    console.print(f"\n[yellow]Total tokens: {total_tokens}/{max_tokens} ({(total_tokens/max_tokens)*100:.1f}%)[/yellow]")
                    
                    if chat_state.summary:
                        summary_tokens = count_tokens(chat_state.summary, model)
                        console.print(f"[yellow]Summary tokens: {summary_tokens} (saves {total_tokens - summary_tokens} tokens if used)[/yellow]")
                    continue
                
                elif command == '/kb' and VECTOR_DB_AVAILABLE:
                    if not kb_instance:
                        console.print("[red]Knowledge base is not available. Please check your installation.[/red]")
                        continue
                        
                    kb_args = args.strip().split(' ', 1)
                    kb_command = kb_args[0].lower() if kb_args else ""
                    kb_content = kb_args[1] if len(kb_args) > 1 else ""
                    
                    if kb_command == "status":
                        # Show knowledge base status
                        if not chat_state.kb_enabled:
                            console.print("[yellow]Knowledge base is currently disabled[/yellow]")
                            console.print("[yellow]Use /kb toggle to enable it[/yellow]")
                            continue
                            
                        stats = kb_instance.get_stats()
                        console.print(f"[green]Knowledge base status:[/green]")
                        console.print(f"[green]• Enabled: {chat_state.kb_enabled}[/green]")
                        console.print(f"[green]• Documents: {stats['count']}[/green]")
                        console.print(f"[green]• Path: {stats['path']}[/green]")
                        console.print(f"[green]• Max results: {kb_instance.max_results}[/green]")
                        console.print(f"[green]• Similarity threshold: {kb_instance.similarity_threshold}[/green]")
                        continue
                    
                    elif kb_command == "add":
                        # Add content to knowledge base
                        if not chat_state.kb_enabled:
                            console.print("[yellow]Knowledge base is currently disabled[/yellow]")
                            console.print("[yellow]Use /kb toggle to enable it[/yellow]")
                            continue
                            
                        if not kb_content:
                            console.print("[red]Please provide text to add to the knowledge base[/red]")
                            continue
                            
                        try:
                            doc_id = kb_instance.add_to_knowledge_base(kb_content)
                            if doc_id:
                                console.print(f"[green]Added to knowledge base with ID: {doc_id}[/green]")
                            else:
                                console.print("[yellow]Text was too short to add to knowledge base[/yellow]")
                        except Exception as e:
                            console.print(f"[red]Error adding to knowledge base: {str(e)}[/red]")
                        continue
                            
                    elif kb_command == "search":
                        # Search knowledge base
                        if not chat_state.kb_enabled:
                            console.print("[yellow]Knowledge base is currently disabled[/yellow]")
                            console.print("[yellow]Use /kb toggle to enable it[/yellow]")
                            continue
                            
                        if not kb_content:
                            console.print("[red]Please provide a search query[/red]")
                            continue
                            
                        try:
                            kb_results = kb_instance.search_knowledge_base(kb_content)
                            if kb_results:
                                console.print(f"[green]Found {len(kb_results)} results:[/green]")
                                for i, result in enumerate(kb_results):
                                    similarity_pct = result["similarity"] * 100
                                    source = result["metadata"].get("source", "unknown")
                                    timestamp = result["metadata"].get("timestamp", "unknown")
                                    
                                    # Format the result
                                    console.print(f"[cyan]Result {i+1} (Similarity: {similarity_pct:.1f}%)[/cyan]")
                                    console.print(f"[dim]Source: {source}, Added: {timestamp}[/dim]")
                                    console.print(Panel(result["text"][:500] + ("..." if len(result["text"]) > 500 else ""), 
                                                      border_style="green"))
                            else:
                                console.print("[yellow]No matching results found in knowledge base[/yellow]")
                        except Exception as e:
                            console.print(f"[red]Error searching knowledge base: {str(e)}[/red]")
                        continue
                
                    elif kb_command == "toggle":
                        # Toggle knowledge base
                        chat_state.kb_enabled = not chat_state.kb_enabled
                        status = "enabled" if chat_state.kb_enabled else "disabled"
                        console.print(f"[green]Knowledge base {status}[/green]")
                        continue
                        
                    elif kb_command == "delete":
                        # Delete all documents from knowledge base
                        if not chat_state.kb_enabled:
                            console.print("[yellow]Knowledge base is currently disabled[/yellow]")
                            console.print("[yellow]Use /kb toggle to enable it[/yellow]")
                            continue
                            
                        # Confirm deletion
                        confirmation = input("Are you sure you want to delete ALL documents from the knowledge base? This cannot be undone. (y/n): ")
                        if confirmation.lower() != "y":
                            console.print("[yellow]Operation cancelled[/yellow]")
                            continue
                            
                        try:
                            success = kb_instance.reset_knowledge_base()
                            if success:
                                console.print("[green]Knowledge base has been reset. All documents have been deleted.[/green]")
                            else:
                                console.print("[red]Error resetting knowledge base[/red]")
                        except Exception as e:
                            console.print(f"[red]Error resetting knowledge base: {str(e)}[/red]")
                        continue
                    
                    else:
                        # Show help for knowledge base commands
                        console.print("[cyan]Knowledge Base Commands:[/cyan]")
                        console.print("[green]/kb status[/green] - Show knowledge base status")
                        console.print("[green]/kb add [text][/green] - Add text to knowledge base")
                        console.print("[green]/kb search [query][/green] - Search knowledge base")
                        console.print("[green]/kb toggle[/green] - Enable/disable knowledge base")
                        console.print("[green]/kb delete[/green] - Delete ALL documents from knowledge base")
                        continue
                
                elif command == '/kb' and not VECTOR_DB_AVAILABLE:
                    console.print("[red]Knowledge base feature is not available.[/red]")
                    console.print("[red]Install required dependencies with: pip install chromadb sentence-transformers[/red]")
                    continue
                
                elif (command == '/finetune' or command == '/ft') and FINETUNE_AVAILABLE:
                    # Initialize fine-tuning manager
                    ft_manager = FineTuningManager()
                    
                    if not args:
                        console.print("[yellow]Fine-tuning Commands:[/yellow]")
                        console.print("[green]/finetune status[/green] - Show fine-tuning status and hardware detection")
                        console.print("[green]/finetune create [name] [base_model][/green] - Create a new fine-tuning job")
                        console.print("[green]/finetune dataset [path][/green] - Prepare a dataset for fine-tuning")
                        console.print("[green]/finetune start [name][/green] - Start a fine-tuning job")
                        console.print("[green]/finetune list[/green] - List all fine-tuning jobs")
                        console.print("[green]/finetune export [name][/green] - Export model to Ollama")
                        console.print("[green]/finetune install[/green] - Install fine-tuning dependencies")
                        console.print("[green]/finetune pause [name][/green] - Pause a running fine-tuning job")
                        console.print("[green]/finetune resume [name][/green] - Resume a paused fine-tuning job")
                        console.print("[green]/finetune delete [name][/green] - Delete a fine-tuning job")
                        console.print("[green]/finetune models[/green] - List available Ollama models for fine-tuning")
                        console.print("[green]/finetune progress [name][/green] - Show progress of a fine-tuning job")
                        console.print("[green]/finetune datasets[/green] - List available datasets")
                        console.print("[green]/finetune dataset-set [job_name] [dataset_id][/green] - Update dataset for a job")
                        console.print("[green]/finetune dataset-remove [dataset_id] [--force][/green] - Remove a dataset")
                        console.print("[green]/finetune reset [name][/green] - Reset a job to created state")
                        continue
                    
                    subcmd = args.split()[0] if " " in args else args
                    subcmd_args = args[len(subcmd):].strip() if " " in args else ""
                    
                    if subcmd == "status":
                        # Show hardware detection and fine-tuning status
                        ft_manager.display_status()
                    
                    elif subcmd == "install":
                        # Install dependencies
                        console.print("[yellow]Installing fine-tuning dependencies...[/yellow]")
                        if ft_manager.install_dependencies():
                            console.print("[green]Dependencies installed successfully[/green]")
                        else:
                            console.print("[red]Failed to install dependencies[/red]")
                    
                    elif subcmd == "create":
                        # Create a new fine-tuning job
                        create_args = subcmd_args.split()
                        if len(create_args) < 2:
                            console.print("[red]Error: Missing arguments[/red]")
                            console.print("[yellow]Usage: /finetune create [name] [base_model] [param=value] ...[/yellow]")
                            console.print("[yellow]Example: /finetune create my_job llama2 batch_size=16[/yellow]")
                            continue
                        
                        name = create_args[0]
                        base_model = create_args[1]
                        
                        # Parse additional parameters
                        params = {}
                        for arg in create_args[2:]:
                            if "=" in arg:
                                key, value = arg.split("=", 1)
                                try:
                                    # Try to convert to appropriate type
                                    if value.lower() == "true":
                                        params[key] = True
                                    elif value.lower() == "false":
                                        params[key] = False
                                    elif "." in value:
                                        params[key] = float(value)
                                    else:
                                        params[key] = int(value)
                                except ValueError:
                                    params[key] = value
                        
                        try:
                            if ft_manager.create_job(name, base_model, parameters=params):
                                console.print(f"[green]Created fine-tuning job:[/green] {name}")
                                console.print(f"[green]Base model:[/green] {base_model}")
                                job_info = ft_manager.get_job(name)
                                if job_info:
                                    console.print(f"[green]Framework:[/green] {job_info.get('framework', 'unknown')}")
                                    console.print(f"[green]Parameters:[/green]")
                                    for key, value in job_info.get("parameters", {}).items():
                                        console.print(f"  [green]{key}:[/green] {value}")
                            else:
                                console.print(f"[red]Failed to create job: {name}[/red]")
                        except Exception as e:
                            console.print(f"[red]Error creating job: {str(e)}[/red]")
                    
                    elif subcmd == "list":
                        # List all fine-tuning jobs
                        jobs = ft_manager.get_jobs()
                        if not jobs:
                            console.print("[yellow]No fine-tuning jobs found[/yellow]")
                        else:
                            console.print("[green]Fine-tuning jobs:[/green]")
                            for job_name, job_info in jobs.items():
                                status = job_info.get('status', 'unknown')
                                base_model = job_info.get('base_model', 'unknown')
                                
                                # Get status color
                                if status == "running":
                                    status_color = "green"
                                elif status == "completed":
                                    status_color = "blue"
                                elif status == "paused":
                                    status_color = "yellow"
                                elif status == "failed":
                                    status_color = "red"
                                else:
                                    status_color = "white"
                                
                                console.print(f"  [bold]{job_name}[/bold] - Status: [{status_color}]{status}[/{status_color}] - Model: {base_model}")
                    
                    elif subcmd == "export":
                        # Export model to Ollama
                        export_args = subcmd_args.split()
                        if not export_args:
                            console.print("[red]Error: Missing name argument[/red]")
                            console.print("[yellow]Usage: /finetune export [name] [--target target_name][/yellow]")
                            continue
                        
                        name = export_args[0]
                        target_name = None
                        
                        # Check for target name
                        if len(export_args) > 2 and export_args[1] == "--target":
                            target_name = export_args[2]
                        
                        try:
                            if ft_manager.export_job(name, target_name):
                                console.print(f"[green]Model exported to Ollama: {target_name or name}[/green]")
                            else:
                                console.print(f"[red]Failed to export model to Ollama: {name}[/red]")
                        except Exception as e:
                            console.print(f"[red]Error exporting model: {str(e)}[/red]")
                    
                    elif subcmd == "models":
                        # List available Ollama models
                        console.print("[yellow]Checking available Ollama models...[/yellow]")
                        models = ft_manager.get_ollama_models()
                        
                        if models:
                            console.print("[green]Available Ollama models for fine-tuning:[/green]")
                            for model in models:
                                console.print(f"  - {model.get('name', 'unknown')} ({model.get('size', 'unknown')})")
                            console.print("\n[yellow]Use one of these models with the create command:[/yellow]")
                            console.print("[yellow]Example: /finetune create my_job llama3[/yellow]")
                        else:
                            console.print("[yellow]No Ollama models found. Make sure you've pulled at least one model with 'ollama pull <model>'.[/yellow]")
                            console.print("[yellow]Example: ollama pull llama3[/yellow]")
                    
                    elif subcmd == "pause":
                        # Pause a fine-tuning job
                        if not subcmd_args:
                            console.print("[red]Error: Missing name argument[/red]")
                            console.print("[yellow]Usage: /finetune pause [name][/yellow]")
                            continue
                        
                        name = subcmd_args
                        if ft_manager.pause_job(name):
                            console.print(f"[green]Paused fine-tuning job: {name}[/green]")
                        else:
                            console.print(f"[red]Failed to pause fine-tuning job: {name}[/red]")
                    
                    elif subcmd == "resume":
                        # Resume a fine-tuning job
                        if not subcmd_args:
                            console.print("[red]Error: Missing name argument[/red]")
                            console.print("[yellow]Usage: /finetune resume [name][/yellow]")
                            continue
                        
                        name = subcmd_args
                        if ft_manager.resume_job(name):
                            console.print(f"[green]Resumed fine-tuning job: {name}[/green]")
                        else:
                            console.print(f"[red]Failed to resume fine-tuning job: {name}[/red]")
                    
                    elif subcmd == "delete":
                        # Delete a fine-tuning job
                        if not subcmd_args:
                            console.print("[red]Error: Missing name argument[/red]")
                            console.print("[yellow]Usage: /finetune delete [name][/yellow]")
                            continue
                        
                        name = subcmd_args
                        if ft_manager.delete_job(name):
                            console.print(f"[green]Deleted fine-tuning job: {name}[/green]")
                        else:
                            console.print(f"[red]Failed to delete fine-tuning job: {name}[/red]")
                    
                    elif subcmd == "start":
                        if not subcmd_args:
                            console.print("[red]Error: Missing name argument[/red]")
                            console.print("[yellow]Usage: /finetune start [name][/yellow]")
                            console.print("[yellow]Example: /finetune start my_job[/yellow]")
                            continue
                        
                        # Start a fine-tuning job
                        name = subcmd_args
                        try:
                            console.print(f"[yellow]Starting fine-tuning job: {name}[/yellow]")
                            if ft_manager.start_job(name):
                                console.print(f"[green]Fine-tuning job started: {name}[/green]")
                                console.print("[yellow]Use /finetune progress to check the status of the job[/yellow]")
                            else:
                                console.print(f"[red]Failed to start fine-tuning job: {name}[/red]")
                        except Exception as e:
                            console.print(f"[red]Error starting job: {str(e)}[/red]")
                    
                    elif subcmd == "progress":
                        # Show progress of a fine-tuning job
                        if not subcmd_args:
                            console.print("[red]Error: Missing name argument[/red]")
                            console.print("[yellow]Usage: /finetune progress [name][/yellow]")
                            continue
                        
                        name = subcmd_args
                        ft_manager.display_job_progress(name)
                    
                    elif subcmd == "dataset":
                        if not subcmd_args:
                            console.print("[red]Error: Missing path argument[/red]")
                            console.print("[yellow]Usage: /finetune dataset [path][/yellow]")
                            console.print("[yellow]Example: /finetune dataset ./my_dataset.json[/yellow]")
                            console.print("[yellow]You can also drag and drop a file into the terminal.[/yellow]")
                            continue
                            
                        # Prepare a dataset for fine-tuning
                        path = subcmd_args.strip()
                        
                        # Handle paths with quotes (from drag and drop)
                        if (path.startswith('"') and path.endswith('"')) or (path.startswith("'") and path.endswith("'")):
                            path = path[1:-1]
                        
                        try:
                            dataset_id = ft_manager.prepare_dataset(path)
                            if dataset_id:
                                console.print(f"[green]Dataset prepared: {dataset_id}[/green]")
                            else:
                                console.print(f"[red]Failed to prepare dataset from: {path}[/red]")
                        except Exception as e:
                            console.print(f"[red]Error preparing dataset: {str(e)}[/red]")
                    
                    elif subcmd == "datasets":
                        # List available datasets
                        datasets = ft_manager.get_datasets()
                        if not datasets:
                            console.print("[yellow]No datasets available. Use /finetune dataset to prepare a dataset.[/yellow]")
                            continue
                        
                        console.print("[yellow]Available datasets:[/yellow]")
                        from rich.table import Table
                        table = Table(show_header=True)
                        table.add_column("ID", style="green")
                        table.add_column("Original Path", style="blue")
                        table.add_column("Created", style="yellow")
                        
                        for dataset_id, dataset_info in datasets.items():
                            # Handle the case where created_at might be a string or not present
                            created_at_value = dataset_info.get("created_at", 0)
                            try:
                                # Convert to float first to handle both int and string representations
                                created_at = datetime.datetime.fromtimestamp(float(created_at_value)).strftime("%Y-%m-%d %H:%M:%S")
                            except (ValueError, TypeError):
                                created_at = "Unknown"
                            table.add_row(
                                dataset_id,
                                dataset_info.get("original_path", "Unknown"),
                                created_at
                            )
                        
                        console.print(table)
                    
                    elif subcmd == "dataset-set":
                        # Update dataset for a job
                        args_parts = subcmd_args.split()
                        if len(args_parts) < 2:
                            console.print("[red]Error: Missing arguments[/red]")
                            console.print("[yellow]Usage: /finetune dataset-set [job_name] [dataset_id][/yellow]")
                            console.print("[yellow]Example: /finetune dataset-set my_job dataset_123[/yellow]")
                            console.print("[yellow]Use /finetune datasets to see available datasets[/yellow]")
                            continue
                        
                        job_name = args_parts[0]
                        dataset_id = args_parts[1]
                        
                        if ft_manager.update_job_dataset(job_name, dataset_id):
                            console.print(f"[green]Dataset for job {job_name} updated to {dataset_id}[/green]")
                        else:
                            console.print(f"[red]Failed to update dataset for job {job_name}[/red]")
                    
                    elif subcmd == "dataset-remove":
                        # Remove a dataset
                        args_parts = subcmd_args.split()
                        if not args_parts:
                            console.print("[red]Error: Missing dataset ID[/red]")
                            console.print("[yellow]Usage: /finetune dataset-remove [dataset_id] [--force][/yellow]")
                            console.print("[yellow]Example: /finetune dataset-remove dataset_123[/yellow]")
                            console.print("[yellow]Use --force to remove even if used by jobs[/yellow]")
                            continue
                        
                        dataset_id = args_parts[0]
                        force = "--force" in args_parts
                        
                        if ft_manager.remove_dataset(dataset_id, force):
                            console.print(f"[green]Dataset {dataset_id} removed successfully[/green]")
                        else:
                            console.print(f"[red]Failed to remove dataset {dataset_id}[/red]")
                    
                    elif subcmd == "reset":
                        # Reset a job to created state
                        if not subcmd_args:
                            console.print("[red]Error: Missing job name[/red]")
                            console.print("[yellow]Usage: /finetune reset [name][/yellow]")
                            console.print("[yellow]Example: /finetune reset my_job[/yellow]")
                            continue
                        
                        job_name = subcmd_args
                        if ft_manager.reset_job(job_name):
                            console.print(f"[green]Successfully reset job {job_name} to 'created' state[/green]")
                            console.print(f"[green]You can now start the job again with /finetune start {job_name}[/green]")
                        else:
                            console.print(f"[red]Failed to reset job {job_name}[/red]")
                    
                    else:
                        console.print("[red]Invalid fine-tuning command[/red]")
                        console.print("[yellow]Try /finetune without arguments to see available commands[/yellow]")
                    
                    continue
                
                elif (command == '/finetune' or command == '/ft') and not FINETUNE_AVAILABLE:
                    console.print("[red]Fine-tuning feature is not available.[/red]")
                    console.print("[red]Make sure the finetune.py and hardware_detection.py files exist in the same directory.[/red]")
                    continue
                
                elif command == '/create':
                    # Handle file creation command
                    if not args:
                        console.print("[red]Error: Missing arguments for /create command.[/red]")
                        console.print("[yellow]Usage: /create [filename] [content][/yellow]")
                        console.print("[yellow]Example: /create data.csv 'Name,Age,City\\nJohn,30,New York\\nJane,25,Boston'[/yellow]")
                        console.print("[yellow]Natural language: /create a haiku about nature and save to nature.txt[/yellow]")
                        continue
                    
                    # Check if this is a natural language request
                    if ' and save to ' in args.lower() or ' and save as ' in args.lower() or ' save to ' in args.lower() or ' save as ' in args.lower() or ' in ' in args.lower() or ' named ' in args.lower() or ' called ' in args.lower() or ' to ' in args.lower():
                        # Process natural language request
                        console.print("[yellow]Processing natural language file creation request...[/yellow]")
                        
                        # Extract filename from the request using various patterns
                        filename = None
                        
                        # Try different patterns to extract the filename
                        patterns = [
                            r'(?:save to|save as|save it as|save it to) ([a-zA-Z0-9_\-\.~\/]+)',
                            r'(?:in|to|named|called) ([a-zA-Z0-9_\-\.~\/]+\.(?:txt|csv|docx?|xlsx?|pdf))',
                            r'([a-zA-Z0-9_\-\.~\/]+\.(?:txt|csv|docx?|xlsx?|pdf))',
                            r'save (?:it )?(?:as|to) ([a-zA-Z0-9_\-\.~\/]+)'
                        ]
                        
                        for pattern in patterns:
                            filename_match = re.search(pattern, args, re.IGNORECASE)
                            if filename_match:
                                filename = filename_match.group(1).strip()
                                break
                        
                        if not filename:
                            console.print("[red]Error: Could not identify filename in request.[/red]")
                            console.print("[yellow]Please specify a filename, e.g., 'save to myfile.txt'[/yellow]")
                            continue
                        
                        # Remove the filename part from the request
                        content_request = args
                        for pattern in [
                            r'(?:and)?\s*(?:save to|save as|save it as|save it to) [a-zA-Z0-9_\-\.~\/]+',
                            r'(?:in|to|named|called) [a-zA-Z0-9_\-\.~\/]+\.(?:txt|csv|docx?|xlsx?|pdf)',
                            r'(?:in|to|named|called) [a-zA-Z0-9_\-\.~\/]+',
                            r'save (?:it )?(?:as|to) [a-zA-Z0-9_\-\.~\/]+'
                        ]:
                            content_request = re.sub(pattern, '', content_request, flags=re.IGNORECASE).strip()
                        
                        # Generate content using the model
                        console.print(f"[yellow]Generating content for: {content_request}[/yellow]")
                        console.print(f"[yellow]Creating file: {filename}[/yellow]")
                        
                        # Prepare messages for the model
                        messages = [
                            {"role": "system", "content": "You are a helpful assistant that generates content based on user requests. Be concise and creative."},
                            {"role": "user", "content": content_request}
                        ]
                        
                        # Generate content using the model
                        try:
                            response = generate_chat_completion(model, messages)
                            content = response.strip()
                            
                            # Expand ~ in filename to user's home directory
                            filename = os.path.expanduser(filename)
                            
                            # Create the file
                            console.print(f"[yellow]Creating file: {filename}[/yellow]")
                            console.print(f"[cyan]Generated content:[/cyan]\n{content}\n")
                            success, message = create_file(content, filename)
                            
                            if success:
                                console.print(f"[green]{message}[/green]")
                            else:
                                console.print(f"[red]{message}[/red]")
                        except Exception as e:
                            console.print(f"[red]Error generating content: {str(e)}[/red]")
                    else:
                        # Traditional format: filename content
                        try:
                            # Split only on the first space to get filename and content
                            parts = args.split(' ', 1)
                            if len(parts) < 2:
                                console.print("[red]Error: Invalid format. Please provide both filename and content.[/red]")
                                console.print("[yellow]Usage: /create [filename] [content][/yellow]")
                                continue
                                
                            filename = parts[0]
                            content = parts[1]
                        except ValueError:
                            console.print("[red]Error: Invalid format. Please provide both filename and content.[/red]")
                            console.print("[yellow]Usage: /create [filename] [content][/yellow]")
                            continue
                        
                        # Expand ~ in filename to user's home directory
                        filename = os.path.expanduser(filename)
                        
                        # Create the file
                        console.print(f"[yellow]Creating file: {filename}[/yellow]")
                        success, message = create_file(content, filename)
                        
                        if success:
                            console.print(f"[green]{message}[/green]")
                        else:
                            console.print(f"[red]{message}[/red]")
                    
                    continue
                
                # Filesystem commands
                elif command == 'fs':
                    if not FILESYSTEM_AVAILABLE:
                        console.print("[red]Filesystem functionality not available. Install the filesystem integration module first.[/red]")
                        console.print("[yellow]Run: python install_filesystem_mcp.py[/yellow]")
                        continue
                    
                    # Handle filesystem commands
                    response = handle_fs_command(args.split())
                    console.print(response)
                    continue
                
                # Natural language filesystem commands
                elif command == '/fsnl':
                    if not FILESYSTEM_MCP_AVAILABLE:
                        console.print("[red]Filesystem MCP Protocol integration not available. Install the Filesystem MCP Protocol integration first.[/red]")
                        console.print("[yellow]Run: python install_filesystem_mcp_protocol.py[/yellow]")
                        continue
                    
                    # Use the current model from the chat session
                    current_model = model
                    
                    # Get the natural language command
                    nl_command = args if args else Prompt.ask("[cyan]Enter your natural language filesystem command[/cyan]")
                    
                    # Show a spinner while processing
                    with Progress(
                        SpinnerColumn(),
                        TextColumn("[cyan]Processing natural language command...[/cyan]"),
                        transient=True
                    ) as progress:
                        progress.add_task("Processing", total=None)
                        
                        # Handle natural language filesystem command
                        response = handle_filesystem_nl_command(nl_command, current_model)
                    
                    # Display the response as markdown
                    console.print(Markdown(response))
                    continue
                
                # Confluence commands
                elif command == '/confluence':
                    if not CONFLUENCE_MCP_AVAILABLE:
                        console.print("[red]Confluence MCP Protocol integration not available. Install the Confluence MCP Protocol integration first.[/red]")
                        continue
                    
                    # Check if the command is to enter Confluence mode
                    if args.strip().lower() == 'mode':
                        # Enter Confluence mode
                        confluence_mode()
                        continue
                    
                    # Check if Confluence is configured
                    is_configured, message = check_confluence_configuration()
                    if not is_configured:
                        console.print(f"[red]Confluence is not configured: {message}[/red]")
                        setup_now = Prompt.ask("Would you like to set up Confluence integration now?", choices=["y", "n"], default="y")
                        if setup_now.lower() == "y":
                            # Get Confluence configuration
                            console.print("\n[cyan]Confluence Configuration[/cyan]")
                            console.print("You'll need your Confluence Cloud URL and API token to proceed.")
                            console.print("You can generate an API token at https://id.atlassian.com/manage-profile/security/api-tokens")
                            
                            confluence_url = Prompt.ask("Enter your Confluence Cloud URL (e.g., https://your-domain.atlassian.net)")
                            confluence_email = Prompt.ask("Enter your Confluence email address")
                            confluence_token = Prompt.ask("Enter your Confluence API token", password=True)
                            
                            # Save configuration
                            if save_confluence_config(confluence_url, confluence_token, confluence_email):
                                console.print("[green]Confluence configuration saved successfully![/green]")
                                is_configured = True
                            else:
                                console.print("[red]Failed to save Confluence configuration.[/red]")
                                continue
                        else:
                            continue
                    
                    # Get the natural language command
                    nl_command = args if args else Prompt.ask("[cyan]Enter your natural language Confluence command[/cyan]")
                    
                    # Show a spinner while processing
                    with Progress(
                        SpinnerColumn(),
                        TextColumn("[cyan]Processing Confluence command...[/cyan]"),
                        transient=True
                    ) as progress:
                        progress.add_task("Processing", total=None)
                        
                        # Handle natural language Confluence command
                        result = handle_confluence_nl_command(nl_command, model=current_model)
                    
                    # Display the result
                    display_confluence_result(result)
                    continue
                
                # If we get here, it's an unrecognized command
                elif user_input.startswith('/'):
                    console.print("[red]Unrecognized command. Type /help for available commands.[/red]")
                    continue
            
            # Check for search command
            if user_input.lower().startswith("search:"):
                search_query = user_input[7:].strip()  # Remove "search:" prefix
                console.print("[yellow]Performing web search and analysis...[/yellow]")
                
                try:
                    # Perform web search and get analysis
                    search_results = process_enhanced_search(search_query, model)
                    
                    # Add search results to chat history
                    chat_history.append({"role": "user", "content": f"Web search query: {search_query}"})
                    chat_history.append({"role": "assistant", "content": search_results})
                    
                    # Display results
                    # Create thread line for search results
                    if len(chat_history) > 2:  # More than system message + 1 message
                        thread_line = "[dim]│[/dim]"
                        console.print(thread_line)
                    
                    console.print(Panel(
                        Markdown(search_results), 
                        title="[bold green]Search Results & Analysis[/bold green]", 
                        border_style="green",
                        padding=(1, 2)
                    ))
                    continue
                except Exception as e:
                    console.print(f"[red]Error during search: {str(e)}[/red]")
                    continue

            # Add user message to chat history
            chat_history.append({"role": "user", "content": user_input})

            # Display user message in a bubble
            user_panel = Panel(
                Markdown(user_input),
                border_style="blue",
                title="[bold blue]You[/bold blue]",
                padding=(1, 2),
                width=min(len(user_input) + 10, console.width - 20),
                expand=True  # Allow panel to expand with content
            )

            # Create thread line if this isn't the first message
            if len(chat_history) > 2:  # More than system message + 1 user message
                thread_line = "[dim]│[/dim]"
                console.print(thread_line)

            # Print user message right-aligned
            console.print(Align.right(user_panel))
            
            # Regular chat - use streaming response with current context
            try:
                # Prepare messages for context
                context_messages = []
                
                # Always include system message
                for msg in chat_history:
                    if msg["role"] == "system":
                        context_messages.append(msg)
                        break
                
                # If we have a summary, use it instead of the full history
                if chat_state.summary and len(chat_history) > 10:  # Only use summary if conversation is substantial
                    context_messages.append({"role": "system", "content": f"Previous conversation summary: {chat_state.summary}"})
                    
                    # Add the last few messages for immediate context
                    recent_messages = [msg for i, msg in enumerate(chat_history) if i > 0 and i >= len(chat_history) - 5]
                    context_messages.extend(recent_messages)
                else:
                    # Add all messages that aren't excluded, or are pinned
                    for i, msg in enumerate(chat_history):
                        if i > 0:  # Skip system message (already added)
                            if i in chat_state.pinned_messages or i not in chat_state.excluded_messages:
                                context_messages.append(msg)
                
                # Add the current user message
                if not context_messages[-1]["role"] == "user" or context_messages[-1]["content"] != user_input:
                    context_messages.append({"role": "user", "content": user_input})
                
                # Search knowledge base for relevant information if enabled
                kb_context = ""
                if VECTOR_DB_AVAILABLE and kb_instance and chat_state.kb_enabled:
                    try:
                        kb_results = kb_instance.search_knowledge_base(user_input)
                        if kb_results:
                            kb_context = "\n\nRelevant information from knowledge base:\n"
                            for i, result in enumerate(kb_results):
                                similarity_pct = result["similarity"] * 100
                                source = result["metadata"].get("source", "unknown")
                                timestamp = result["metadata"].get("timestamp", "unknown")
                                
                                # Format the result with thread line and improved styling
                                thread_line = "[dim]│[/dim]"
                                console.print(thread_line)
                                
                                # Create header with similarity info
                                header = f"[bold cyan]Result {i+1} (Similarity: {similarity_pct:.1f}%)[/bold cyan]\n[dim]Source: {source}, Added: {timestamp}[/dim]"
                                
                                # Create panel with result text
                                console.print(Panel(
                                    Markdown(result["text"][:500] + ("..." if len(result["text"]) > 500 else "")),
                                    title=header,
                                    border_style="cyan",
                                    padding=(1, 2)
                                ))
                            kb_context += "\n".join([result["text"] for result in kb_results])
                        else:
                            console.print("[yellow]No matching results found in knowledge base[/yellow]")
                    except Exception as e:
                        config = load_config()
                        if config.get("verbose", False):
                            console.print(f"[yellow]Knowledge base search error: {str(e)}[/yellow]")
                
                # Calculate token count for logging
                current_tokens = sum(count_tokens(msg["content"], model) for msg in context_messages)
                token_display = format_token_count(current_tokens, max_tokens)
                
                # Log token usage if verbose
                config = load_config()
                if config.get("verbose", False):
                    console.print(f"[dim]{token_display}[/dim]")
                
                # Stream the response
                console.print(assistant_separator)
                
                # Create thread line if this isn't the first message
                if len(chat_history) > 2:  # More than system message + 1 user message
                    thread_line = "[dim]│[/dim]"
                    console.print(thread_line)
                
                console.print("\n[green]Assistant:[/green]")
                
                # Use the context_messages instead of just the user input
                # Send the message with the prepared context
                with console.status("[cyan]Thinking...[/cyan]", spinner="dots"):
                    response = send_message(model, "", None, True, context_messages)
                
                # Initialize variables for streaming
                full_response = ""
                buffer = ""
                code_block = False
                code_lang = ""
                code_content = ""
                in_code_block = False
                markdown_buffer = ""
                
                # Create a Live display for updating the markdown content
                from rich.live import Live
                
                # Process the streaming response in real-time
                with Live("", refresh_per_second=10, transient=True) as live:
                    for line in response.iter_lines():
                        if line:
                            try:
                                chunk = json.loads(line)
                                if "message" in chunk and "content" in chunk["message"]:
                                    content = chunk["message"]["content"]
                                    
                                    # Add to the full response
                                    full_response += content
                                    
                                    # Check for code block markers
                                    if "```" in content:
                                        # Handle code block boundaries
                                        parts = content.split("```")
                                        for i, part in enumerate(parts):
                                            if i % 2 == 0:  # Outside code block
                                                if in_code_block:
                                                    code_content += part
                                                else:
                                                    markdown_buffer += part
                                            else:  # Inside code block or language specifier
                                                if not in_code_block:
                                                    in_code_block = True
                                                    # Check if this contains a language specifier
                                                    if part.strip() and not part.strip()[0].isspace():
                                                        code_lang = part.strip().split()[0]
                                                        code_content = part[len(code_lang):].lstrip()
                                                    else:
                                                        code_lang = ""
                                                        code_content = part
                                                else:
                                                    code_content += part
                                                    markdown_buffer += f"```{code_lang}\n{code_content}```"
                                                    code_content = ""
                                                    in_code_block = False
                                    else:
                                        # Regular content (not a code block boundary)
                                        if in_code_block:
                                            code_content += content
                                        else:
                                            markdown_buffer += content
                                    
                                    # Update the live display with the current markdown content
                                    live.update(Markdown(full_response))
                            except json.JSONDecodeError:
                                pass
                
                # Create a panel for the final formatted response
                assistant_panel = Panel(
                    Markdown(full_response),
                    border_style="green", 
                    padding=(1, 2),
                    width=None,
                    expand=True
                )
                
                # Print the final formatted response in a panel
                console.print(assistant_panel)
                
                # Store the response for copying
                last_response = full_response
                
                # Add the response to chat history
                chat_history.append({"role": "assistant", "content": full_response})
                
                # Display token usage
                response_tokens = count_tokens(full_response, model)
                total_tokens = current_tokens + response_tokens
                
                # Create thread line for token info
                thread_line = "[dim]│[/dim]"
                console.print(thread_line)
                
                console.print(Panel(
                    format_token_count(total_tokens, max_tokens),
                    title="[blue]Context Window[/blue]",
                    border_style="blue",
                    padding=(0, 1)
                ))
                
                # Save history if enabled
                save_history(model, chat_history)
                
            except Exception as e:
                console.print(f"\n[red]Error: {str(e)}[/red]")
        
        except KeyboardInterrupt:
            console.print("\n[yellow]Exiting chat...[/yellow]")
            break
        except requests.exceptions.ConnectionError:
            console.print("\n[red]Error: Could not connect to Ollama. Is it running?[/red]")
            break
        except Exception as e:
            console.print(f"\n[red]Error: {str(e)}[/red]")
            config = load_config()
            if config.get("verbose", False):
                import traceback
                console.print(f"\n[red]Traceback: {traceback.format_exc()}[/red]")
            continue

@app.command()
def help():
    """Show help information about available commands"""
    console.clear()
    display_banner()
    
    console.print("\n[bold cyan]Ollama Shell Help[/bold cyan]")
    
    # Basic commands
    console.print("\n[yellow]Basic Commands:[/yellow]")
    console.print("  [green]chat[/green] - Start an interactive chat session with a model")
    console.print("  [green]models[/green] - List all available models")
    console.print("  [green]pull[/green] - Download a new model")
    console.print("  [green]delete[/green] - Delete an installed model")
    console.print("  [green]history[/green] - View chat history")
    console.print("  [green]settings[/green] - View or modify configuration settings")
    console.print("  [green]prompts[/green] - Manage stored system prompts")
    console.print("  [green]analyze[/green] - Analyze an image using vision model")
    
    # Chat features
    console.print("\n[yellow]Chat Features:[/yellow]")
    console.print("  [green]search: [query][/green] - Perform a web search and analyze results")
    console.print("  [green]Ctrl+V[/green] - Toggle drag & drop mode for file sharing")
    console.print("  [green]Ctrl+C[/green] - Copy assistant's response to clipboard")
    
    # Context Management
    console.print("\n[yellow]Context Management Commands:[/yellow]")
    console.print("  [green]/pin [message_number][/green] - Pin a message to keep it in context")
    console.print("  [green]/unpin [message_number][/green] - Unpin a previously pinned message")
    console.print("  [green]/exclude [message_number][/green] - Exclude a message from context")
    console.print("  [green]/include [message_number][/green] - Include a previously excluded message")
    console.print("  [green]/summarize[/green] - Summarize the conversation to save tokens")
    console.print("  [green]/context[/green] - Show current context management status")
    console.print("  [green]/tokens[/green] - Show token usage information")
    console.print("  [green]/help[/green] - Show context management help")
    
    # Knowledge Base
    console.print("\n[yellow]Knowledge Base Commands:[/yellow]")
    console.print("  [green]/kb status[/green] - Show knowledge base status")
    console.print("  [green]/kb add [text][/green] - Add text to knowledge base")
    console.print("  [green]/kb search [query][/green] - Search knowledge base")
    console.print("  [green]/kb toggle[/green] - Enable/disable knowledge base")
    
    # Confluence Integration
    if CONFLUENCE_MCP_AVAILABLE:
        console.print("\n[yellow]Confluence Commands:[/yellow]")
        console.print("  [green]/confluence [command][/green] - Execute Confluence operations using natural language")
        console.print("  [green]/confluence mode[/green] - Enter interactive Confluence mode")
    
    if FINETUNE_AVAILABLE:
        console.print("\n[bold cyan]Fine-Tuning Commands:[/bold cyan]")
        console.print("[green]/finetune status[/green] - Show fine-tuning status and hardware detection")
        console.print("[green]/finetune create [name] [base_model][/green] - Create a new fine-tuning job")
        console.print("[green]/finetune dataset [path][/green] - Prepare a dataset for fine-tuning")
        console.print("[green]/finetune start [name][/green] - Start a fine-tuning job")
        console.print("[green]/finetune list[/green] - List all fine-tuning jobs")
        console.print("[green]/finetune export [name][/green] - Export model to Ollama")
        console.print("[green]/finetune install[/green] - Install fine-tuning dependencies")
        console.print("[green]/finetune pause [name][/green] - Pause a running fine-tuning job")
        console.print("[green]/finetune resume [name][/green] - Resume a paused fine-tuning job")
        console.print("[green]/finetune delete [name][/green] - Delete a fine-tuning job")
        console.print("[green]/finetune models[/green] - List available Ollama models for fine-tuning")
        console.print("[green]/finetune progress [name][/green] - Show progress of a fine-tuning job")
    
    # File handling
    console.print("\n[yellow]File Handling:[/yellow]")
    console.print("  Drag & drop files into the chat for analysis")
    
    # Filesystem commands
    if FILESYSTEM_AVAILABLE:
        console.print("\n[bold cyan]Filesystem Commands:[/bold cyan]")
        console.print("  [green]/fs help[/green] - Show filesystem commands help")
        console.print("  [green]/fs ls [path][/green] - List directory contents")
        console.print("  [green]/fs read <path>[/green] - Read file content")
        console.print("  [green]/fs write <path> <content>[/green] - Write content to a file")
        console.print("  [green]/fs mkdir <path>[/green] - Create a directory")
        console.print("  [green]/fs analyze <path>[/green] - Analyze text file properties")
        console.print("  [green]/fs zip <output> <sources...>[/green] - Create a ZIP archive")
        console.print("  [green]/fs unzip <zip> <output>[/green] - Extract a ZIP archive")
    # Natural language filesystem commands
    if FILESYSTEM_MCP_AVAILABLE:
        console.print("\n[bold cyan]Natural Language Filesystem Commands:[/bold cyan]")
        console.print("  [green]/fsnl [command][/green] - Execute filesystem operations using natural language")
        console.print("  [yellow]Examples:[/yellow]")
        console.print("    [green]/fsnl list all files in my documents folder[/green]")
        console.print("    [green]/fsnl create a text file named notes.txt with my meeting agenda[/green]")
        console.print("    [green]/fsnl find all images in my downloads folder[/green]")
    
    console.print("  Supported formats: PDF, DOCX, TXT, code files, images (with vision models)")
    console.print("  [green]/create [filename] [content][/green] - Create a file with specified content")
    console.print("  [green]/create [request] and save to [filename][/green] - Generate and save content")
    console.print("  Supported file types: TXT, CSV, DOC/DOCX, XLS/XLSX, PDF")
    
    # Press Enter to continue
    Prompt.ask("\n[cyan]Press Enter to continue[/cyan]")

@app.command()
def chat(
    model: Optional[str] = typer.Option(None, help="Model to use for chat"),
    system_prompt: Optional[str] = typer.Option(None, help="System prompt to use"),
    context_files: Optional[list[str]] = typer.Option(None, help="List of files to include in context")
):
    """Start an interactive chat session with the specified model"""
    try:
        # If no model specified, use default from config
        if model is None:
            config = load_config()
            model = config.get("default_model", "llama2")
        
        # Convert model from OptionInfo to string if needed
        if not isinstance(model, str):
            model = str(model)
        
        # Start interactive chat
        interactive_chat(model, system_prompt, context_files)
    except Exception as e:
        console.print(f"[red]Error starting chat: {str(e)}[/red]")

@app.command()
def models():
    """List available models and their details"""
    try:
        response = requests.get(f"{OLLAMA_API}/tags")
        if response.status_code == 200:
            models = response.json()["models"]
            
            # Create a table of models
            table = Table(title="Available Models")
            table.add_column("Name", style="cyan")
            table.add_column("Size", style="green")
            table.add_column("Modified", style="yellow")
            
            for model in models:
                size = f"{model['size'] / 1024 / 1024 / 1024:.1f}GB"
                table.add_row(
                    model['name'],
                    size,
                    model.get('modified', 'N/A')
                )
            
            display_banner()
            console.print(table)
        else:
            console.print("[red]Error fetching models[/red]")
    except Exception as e:
        console.print(f"[red]Error: {str(e)}[/red]")

@app.command()
def pull(
    model: str = typer.Argument(..., help="Name of the model to download")
):
    """Download a new model from Ollama"""
    try:
        console.print(f"\n[cyan]Downloading model: {model}[/cyan]")
        
        # Track progress
        seen_statuses = set()
        current_status = None
        download_started = False
        last_update = time.time()
        timeout = 30  # seconds
        
        with console.status("[cyan]Starting download...[/cyan]") as status:
            response = requests.post(
                f"{OLLAMA_API}/pull",
                json={"name": model},
                stream=True,
                timeout=10  # Initial connection timeout
            )
            
            if response.status_code == 200:
                for line in response.iter_lines():
                    if line:
                        try:
                            data = json.loads(line)
                            
                            # Reset timeout counter on any data
                            last_update = time.time()
                            
                            if 'status' in data:
                                current_status = data['status']
                                
                                # Show progress for downloading
                                if 'downloading' in current_status.lower():
                                    download_started = True
                                    if 'completed' in data:
                                        completed = data['completed']
                                        total = data.get('total', 0)
                                        if total > 0:
                                            percentage = (completed / total) * 100
                                            status.update(f"[cyan]Downloading: {percentage:.1f}% ({completed}/{total} chunks)[/cyan]")
                                        else:
                                            status.update(f"[cyan]Downloading: {completed} chunks completed[/cyan]")
                                
                                # Show other status updates
                                elif current_status not in seen_statuses:
                                    seen_statuses.add(current_status)
                                    if "pulling" in current_status.lower():
                                        status.update(f"[cyan]Pulling model files...[/cyan]")
                                    elif "verifying" in current_status.lower():
                                        status.update(f"[yellow]Verifying download...[/yellow]")
                                    elif "writing" in current_status.lower():
                                        status.update(f"[green]Writing model to disk...[/green]")
                                    else:
                                        status.update(f"[cyan]{current_status}[/cyan]")
                            
                            if 'error' in data:
                                console.print(f"\n[red]Error: {data['error']}[/red]")
                                return
                            
                        except json.JSONDecodeError:
                            # Skip invalid JSON lines
                            continue
                            
                        # Check for timeout
                        if time.time() - last_update > timeout:
                            raise TimeoutError("No progress updates received for too long")
                
                if download_started:
                    console.print(f"\n[green]Successfully downloaded {model}![/green]")
                else:
                    console.print(f"\n[green]Model {model} is already up to date![/green]")
            
            else:
                console.print(f"\n[red]Error: Server returned status code {response.status_code}[/red]")
                if response.text:
                    console.print(f"[red]Details: {response.text}[/red]")
                
    except requests.exceptions.Timeout:
        console.print("\n[red]Error: Connection to Ollama timed out. Please check if Ollama is running.[/red]")
    except TimeoutError as e:
        console.print(f"\n[red]Error: {str(e)}. The download may have stalled.[/red]")
    except requests.exceptions.ConnectionError:
        console.print("\n[red]Error: Could not connect to Ollama. Please check if it's running.[/red]")
    except Exception as e:
        console.print(f"\n[red]Error: {str(e)}[/red]")

@app.command()
def delete(
    model: str = typer.Argument(..., help="Name of the model to delete")
):
    """Delete a model from Ollama"""
    try:
        # If no model specified, show list of installed models
        if not model or model == "list":
            available_models = get_available_models()
            if not available_models:
                console.print("[yellow]No models installed. Is Ollama running?[/yellow]")
                return
            
            console.print("\n[cyan]Installed Models:[/cyan]")
            for i, m in enumerate(available_models, 1):
                console.print(f"[green]{i}[/green]. {m}")
            
            choice = Prompt.ask("\nSelect model to delete (number or name)", default="")
            if not choice:
                console.print("[yellow]Operation cancelled.[/yellow]")
                return
            
            # Handle numeric choice
            if choice.isdigit() and 1 <= int(choice) <= len(available_models):
                model = available_models[int(choice) - 1]
            else:
                model = choice
        
        # Confirm deletion
        if not Prompt.ask(f"\n[yellow]Are you sure you want to delete {model}?[/yellow]", choices=["y", "n", "yes", "no"], default="n").lower() in ['y', 'yes']:
            console.print("[yellow]Operation cancelled.[/yellow]")
            return
        
        with console.status(f"[cyan]Deleting model: {model}...[/cyan]"):
            response = requests.delete(
                f"{OLLAMA_API}/delete",
                json={"name": model}
            )
            
            if response.status_code == 200:
                console.print(f"\n[green]Successfully deleted {model}![/green]")
            else:
                console.print(f"\n[red]Error deleting model: {response.text}[/red]")
    except Exception as e:
        console.print(f"\n[red]Error: {str(e)}[/red]")

@app.command()
def prompt(
    model: str = typer.Option("llama2", help="Model to use"),
    prompt_text: str = typer.Argument(..., help="The prompt to send")
):
    """Send a single prompt to the model"""
    with console.status("[cyan]Thinking...[/cyan]"):
        response = send_message(model, prompt_text)
        if response:
            console.print(Panel(Markdown(response["message"]["content"]), border_style="purple"))

@app.command()
def config_settings(
    default_model: Optional[str] = typer.Option(None, help="Set default model"),
    verbose: Optional[bool] = typer.Option(None, help="Toggle verbose mode"),
    save_history: Optional[bool] = typer.Option(None, help="Toggle history saving"),
    temperature: Optional[float] = typer.Option(None, help="Set temperature (0.0-1.0)"),
    context_length: Optional[int] = typer.Option(None, help="Set context length"),
    show: bool = typer.Option(False, help="Show current configuration")
):
    """View or modify configuration settings"""
    current_config = load_config()
    
    if show:
        table = Table(title="Current Configuration")
        table.add_column("Setting", style="cyan")
        table.add_column("Value", style="green")
        
        for key, value in current_config.items():
            table.add_row(key, str(value))
        
        display_banner()
        console.print(table)
        return
    
    # Update configuration if any options provided
    if any(x is not None for x in [default_model, verbose, save_history, temperature, context_length]):
        if default_model:
            current_config["default_model"] = default_model
        if verbose is not None:
            current_config["verbose"] = verbose
        if save_history is not None:
            current_config["save_history"] = save_history
        if temperature is not None:
            current_config["temperature"] = max(0.0, min(1.0, temperature))
        if context_length is not None:
            current_config["context_length"] = context_length
        
        save_config(current_config)
        console.print("\n[green]Configuration updated:[/green]")
        config_settings(show=True)

def interactive_config():
    """Interactive configuration editor"""
    current_config = load_config()
    
    while True:
        console.clear()
        display_banner()
        
        # Display current configuration
        table = Table(title="Current Configuration")
        table.add_column("Number", style="green")
        table.add_column("Setting", style="cyan")
        table.add_column("Value", style="white")
        
        # Convert config items to list for numbered access
        config_items = list(current_config.items())
        for i, (key, value) in enumerate(config_items, 1):
            table.add_row(f"[green]{i}[/green]", key, str(value))
        
        # Add Confluence configuration option if available
        if CONFLUENCE_MCP_AVAILABLE:
            is_configured, _ = check_confluence_configuration()
            status = "[green]Configured[/green]" if is_configured else "[red]Not configured[/red]"
            table.add_row(f"[green]{len(config_items) + 1}[/green]", "confluence_settings", status)
            
        # Add Jira configuration option if available
        if JIRA_MCP_AVAILABLE:
            is_configured, _ = check_jira_configuration()
            status = "[green]Configured[/green]" if is_configured else "[red]Not configured[/red]"
            jira_option_number = len(config_items) + 1
            if CONFLUENCE_MCP_AVAILABLE:
                jira_option_number += 1
            table.add_row(f"[green]{jira_option_number}[/green]", "jira_settings", status)
        
        console.print(table)
        console.print("\n[cyan]Options:[/cyan]")
        console.print(f"[green]1-{len(config_items)}[/green]: Edit setting")
        if CONFLUENCE_MCP_AVAILABLE:
            console.print(f"[green]{len(config_items) + 1}[/green]: Configure Confluence")
        if JIRA_MCP_AVAILABLE:
            jira_option_number = len(config_items) + 1
            if CONFLUENCE_MCP_AVAILABLE:
                jira_option_number += 1
            console.print(f"[green]{jira_option_number}[/green]: Configure Jira")
        console.print("[green]s[/green]: Save and exit")
        console.print("[green]x[/green]: Exit without saving")
        
        choice = Prompt.ask("\nEnter your choice")
        
        if choice.lower() == 'x':
            break
        elif choice.lower() == 's':
            save_config(current_config)
            break
        elif choice.isdigit():
            choice_num = int(choice)
            if 1 <= choice_num <= len(config_items):
                index = choice_num - 1
                key, current_value = config_items[index]
                
                if key == "default_model":
                    # Show available models for selection
                    response = requests.get(f"{OLLAMA_API}/tags")
                    if response.status_code == 200:
                        models = [model["name"] for model in response.json()["models"]]
                        console.print("\nAvailable models:")
                        for i, model in enumerate(models):
                            console.print(f"[green]{i+1}[/green]. {model}")
                        model_choice = Prompt.ask("\nSelect model number", default="1")
                        if model_choice.isdigit() and 1 <= int(model_choice) <= len(models):
                            current_config[key] = models[int(model_choice) - 1]
            elif CONFLUENCE_MCP_AVAILABLE and choice_num == len(config_items) + 1:
                # Configure Confluence settings
                console.print("\n[cyan]Confluence Configuration[/cyan]")
                console.print("You'll need your Confluence Cloud URL and API token to proceed.")
                console.print("You can generate an API token at https://id.atlassian.com/manage-profile/security/api-tokens")
                
                confluence_url = Prompt.ask("Enter your Confluence Cloud URL (e.g., https://your-domain.atlassian.net)")
                confluence_email = Prompt.ask("Enter your Confluence email address")
                confluence_token = Prompt.ask("Enter your Confluence API token", password=True)
                
                # Save configuration
                if save_confluence_config(confluence_url, confluence_token, confluence_email):
                    console.print("[green]Confluence configuration saved successfully![/green]")
                else:
                    console.print("[red]Failed to save Confluence configuration.[/red]")
            
            elif JIRA_MCP_AVAILABLE and choice_num == (len(config_items) + 1 + (1 if CONFLUENCE_MCP_AVAILABLE else 0)):
                # Configure Jira settings
                console.print("\n[cyan]Jira Configuration[/cyan]")
                console.print("You'll need your Jira URL and API token to proceed.")
                console.print("You can generate an API token at https://id.atlassian.com/manage-profile/security/api-tokens")
                
                jira_url = Prompt.ask("Enter your Jira URL (e.g., https://your-domain.atlassian.net)")
                jira_email = Prompt.ask("Enter your Jira email address")
                jira_token = Prompt.ask("Enter your Jira API token", password=True)
                
                # Save configuration
                if save_jira_config(jira_url, jira_token, jira_email):
                    console.print("[green]Jira configuration saved successfully![/green]")
                else:
                    console.print("[red]Failed to save Jira configuration.[/red]")
            
            elif isinstance(current_value, bool):
                current_config[key] = not current_value
                console.print(f"\n[green]{key} set to {current_config[key]}[/green]")
            
            elif isinstance(current_value, (int, float)):
                if key == "temperature":
                    new_value = Prompt.ask(f"\nEnter new {key} (0.0-1.0)", default=str(current_value))
                    try:
                        current_config[key] = max(0.0, min(1.0, float(new_value)))
                    except ValueError:
                        console.print("[red]Invalid value. Must be between 0.0 and 1.0[/red]")
                else:
                    new_value = Prompt.ask(f"\nEnter new {key}", default=str(current_value))
                    try:
                        current_config[key] = int(new_value) if isinstance(current_value, int) else float(new_value)
                    except ValueError:
                        console.print("[red]Invalid value[/red]")
            
            else:  # string values
                new_value = Prompt.ask(f"\nEnter new {key}", default=str(current_value))
                current_config[key] = new_value
            
            Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")
        Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")

def manage_prompts():
    """Manage stored system prompts"""
    while True:
        console.clear()
        display_banner()
        
        # Display current prompts
        console.print("\n[bold yellow]Stored System Prompts[/bold yellow]")
        config = load_config()
        if not config["stored_prompts"]:
            console.print("[dim]No stored prompts. Add one to get started![/dim]")
        else:
            table = Table(show_header=True)
            table.add_column("Name", style="cyan")
            table.add_column("Title", style="yellow")
            table.add_column("Preview", style="white")
            
            for name, prompt_data in config["stored_prompts"].items():
                # Truncate long prompts for display
                preview = prompt_data["prompt"][:80] + "..." if len(prompt_data["prompt"]) > 80 else prompt_data["prompt"]
                preview = preview.replace("\n", " ")  # Make it single line
                table.add_row(name, prompt_data["title"], preview)
            
            console.print(table)
        
        # Display options
        console.print("\n[bold yellow]Options[/bold yellow]")
        options = [
            ("1", "Add new prompt"),
            ("2", "Delete prompt"),
            ("3", "View full prompt"),
            ("4", "Search prompts"),
            ("5", "Return to main menu")
        ]
        
        for num, desc in options:
            console.print(f"[green]{num}[/green]. {desc}")
        
        choice = Prompt.ask("\nEnter your choice", choices=[str(i) for i in range(1, 6)])
        choice = int(choice)
        
        if choice == 1:
            name = Prompt.ask("\nEnter a name for the prompt (used for selection)")
            if name in config["stored_prompts"]:
                if not Prompt.ask(f"[yellow]Prompt '{name}' already exists. Overwrite?[/yellow]", choices=["y", "n", "yes", "no"], default="n").lower() in ['y', 'yes']:
                    continue
            
            title = Prompt.ask("Enter a descriptive title for the prompt")
            
            console.print("\nEnter the prompt text (press Enter twice to finish):")
            lines = []
            while True:
                line = input()
                if not line and lines and not lines[-1]:
                    break
                lines.append(line)
            
            prompt_text = "\n".join(lines[:-1])  # Remove the last empty line
            config["stored_prompts"][name] = {
                "title": title,
                "prompt": prompt_text
            }
            save_config(config)
            console.print("[green]Prompt saved successfully![/green]")
        
        elif choice == 2:
            if not config["stored_prompts"]:
                console.print("[yellow]No prompts to delete![/yellow]")
                Prompt.ask("\nPress Enter to continue")
                continue
            
            console.print("\n[dim]Available prompts (press Enter to cancel):[/dim]")
            for name, prompt_data in config["stored_prompts"].items():
                console.print(f"  [cyan]{name}[/cyan]: {prompt_data['title']}")
            
            name = Prompt.ask("\nEnter the name of the prompt to delete", default="", show_default=False)
            if not name:
                console.print("[yellow]Operation cancelled.[/yellow]")
                continue
            
            if name not in config["stored_prompts"]:
                console.print(f"[red]Prompt '{name}' not found.[/red]")
                continue
                
            if Prompt.ask(f"[yellow]Are you sure you want to delete '{config['stored_prompts'][name]['title']}'?[/yellow]", choices=["y", "n", "yes", "no"], default="n").lower() in ['y', 'yes']:
                del config["stored_prompts"][name]
                save_config(config)
                console.print("[green]Prompt deleted successfully![/green]")
        
        elif choice == 3:
            if not config["stored_prompts"]:
                console.print("[yellow]No prompts to view![/yellow]")
                Prompt.ask("\nPress Enter to continue")
                continue
            
            name = Prompt.ask("Enter the name of the prompt to view", choices=list(config["stored_prompts"].keys()))
            prompt_data = config["stored_prompts"][name]
            console.print(f"\n[bold cyan]{prompt_data['title']}[/bold cyan]")
            console.print(Panel(prompt_data["prompt"], border_style="cyan"))
            Prompt.ask("\nPress Enter to continue")
        
        elif choice == 4:
            if not config["stored_prompts"]:
                console.print("[yellow]No prompts to search![/yellow]")
                Prompt.ask("\nPress Enter to continue")
                continue
            
            search_term = Prompt.ask("\nEnter search term").lower()
            matches = []
            
            for name, prompt_data in config["stored_prompts"].items():
                if (search_term in name.lower() or 
                    search_term in prompt_data["title"].lower() or 
                    search_term in prompt_data["prompt"].lower()):
                    matches.append((name, prompt_data))
            
            if not matches:
                console.print("[yellow]No matching prompts found.[/yellow]")
            else:
                console.print("\n[bold cyan]Matching Prompts:[/bold cyan]")
                table = Table(show_header=True)
                table.add_column("Name", style="cyan")
                table.add_column("Title", style="yellow")
                table.add_column("Preview", style="white")
                
                for name, prompt_data in matches:
                    preview = prompt_data["prompt"][:80] + "..." if len(prompt_data["prompt"]) > 80 else prompt_data["prompt"]
                    preview = preview.replace("\n", " ")
                    table.add_row(name, prompt_data["title"], preview)
                
                console.print(table)
            
            Prompt.ask("\nPress Enter to continue")
        
        elif choice == 5:
            break
        
        if choice != 5:
            Prompt.ask("\nPress Enter to continue")

def read_file_content(file_path: str) -> tuple[str, str]:
    """Read and process file content based on file type"""
    file_path = os.path.expanduser(file_path)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    try:
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            # For image files, perform analysis
            analysis = analyze_image(file_path)
            return f"Image Analysis:\n{analysis}", "Image file"
            
        elif ext in ['.pdf']:
            # For PDF files, we'll need PyPDF2
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = '\n'.join(page.extract_text() for page in reader.pages)
                    return text, "PDF document"
            except ImportError:
                raise ImportError("PyPDF2 is required for PDF support. Install it with: pip install PyPDF2")
        
        elif ext in ['.docx', '.doc']:
            # For Word documents, we'll need python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                # Extract text from paragraphs and tables
                text_parts = []
                
                # Get text from paragraphs
                for para in doc.paragraphs:
                    text_parts.append(para.text)
                
                # Get text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text.strip())
                        text_parts.append(" | ".join(row_text))
                
                text = "\n".join(text_parts)
                return text, "Word document"
            except ImportError:
                raise ImportError("python-docx is required for Word document support. Install it with: pip install python-docx")
        
        elif ext in ['.xlsx', '.xls']:
            # For Excel files, we'll need pandas and openpyxl
            try:
                import pandas as pd
                
                # Read all sheets
                sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl' if ext == '.xlsx' else 'xlrd')
                
                text_parts = []
                # Process each sheet
                for sheet_name, df in sheets.items():
                    text_parts.append(f"\n--- Sheet: {sheet_name} ---\n")
                    
                    # Convert dataframe to string representation
                    sheet_text = df.to_string(index=False)
                    text_parts.append(sheet_text)
                
                text = "\n".join(text_parts)
                return text, "Excel spreadsheet"
            except ImportError:
                raise ImportError("pandas and openpyxl/xlrd are required for Excel support. Install them with: pip install pandas openpyxl xlrd")
        
        elif ext in ['.md', '.txt', '.py', '.js', '.html', '.css', '.json', '.yaml', '.yml'] or not ext:
            # Text files can be read directly
            # Also handle files without extensions as text files
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                    file_type = f"{ext[1:].upper()} file" if ext else "Text file"
                    return text, file_type
            except UnicodeDecodeError:
                # If UTF-8 fails, try with system default encoding
                with open(file_path, 'r') as file:
                    text = file.read()
                    file_type = f"{ext[1:].upper()} file" if ext else "Text file"
                    return text, file_type
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")
            
    except Exception as e:
        raise Exception(f"Error reading file: {str(e)}")

def prepare_document_context(files: list[str]) -> str:
    """Prepare context from multiple documents"""
    context = []
    for file_path in files:
        try:
            content, file_type = read_file_content(file_path)
            context.append(f"\n--- {file_type}: {os.path.basename(file_path)} ---\n{content}")
        except Exception as e:
            console.print(f"[yellow]Warning: Could not read {file_path}: {str(e)}[/yellow]")
    
    return "\n".join(context)

def display_menu():
    """Display interactive menu and handle command selection"""
    # Define commands with their functions
    commands = {
        "chat": chat,
        "models": models,
        "pull": pull,
        "delete": lambda: delete(None),
        "history": view_history,
        "settings": interactive_config,
        "prompts": manage_prompts,
        "analyze": lambda: analyze_from_menu(),
        "terminal": terminal_mode,
        "help": help,
        "exit": None
    }
    
    # Add filesystem mode if available
    if FILESYSTEM_AVAILABLE:
        commands["filesystem"] = filesystem_mode
    
    # Add Confluence mode if available
    if CONFLUENCE_MCP_AVAILABLE:
        commands["confluence"] = confluence_mode
        
    # Add Jira mode if available
    if JIRA_MCP_AVAILABLE:
        commands["jira"] = jira_mode
    
    # Add fine-tuning command if available
    if FINETUNE_AVAILABLE:
        commands["finetune"] = lambda: interactive_chat(load_config().get("default_model", "llama2"), 
                                                      system_prompt="/finetune status")
    
    # Menu items with numbers
    base_menu_items = [
        ("1", "chat", "Start an interactive chat session with a model"),
        ("2", "models", "List all available models"),
        ("3", "pull", "Download a new model"),
        ("4", "delete", "Delete an installed model"),
        ("5", "history", "View chat history"),
        ("6", "settings", "View or modify configuration settings"),
        ("7", "prompts", "Manage stored system prompts"),
        ("8", "analyze", "Analyze an image using vision model"),
        ("9", "terminal", "Execute system commands in a terminal"),
    ]
    
    # Add filesystem, confluence, and fine-tuning options if available
    menu_items = base_menu_items.copy()
    menu_option_number = 10
    
    # Add filesystem option if available
    if FILESYSTEM_AVAILABLE:
        menu_items.append((str(menu_option_number), "filesystem", "Access filesystem operations"))
        menu_option_number += 1
    
    # Add Confluence option if available
    if CONFLUENCE_MCP_AVAILABLE:
        menu_items.append((str(menu_option_number), "confluence", "Access Confluence knowledge base"))
        menu_option_number += 1
        
    # Add Jira option if available
    if JIRA_MCP_AVAILABLE:
        menu_items.append((str(menu_option_number), "jira", "Access Jira issue tracking"))
        menu_option_number += 1
    
    # Add fine-tuning option if available
    if FINETUNE_AVAILABLE:
        menu_items.append((str(menu_option_number), "finetune", "Fine-tune models with Unsloth or MLX"))
        menu_option_number += 1
    
    # Add help and exit options
    menu_items.append((str(menu_option_number), "help", "Show help information"))
    menu_option_number += 1
    menu_items.append((str(menu_option_number), "exit", "Exit the application"))
    
    while True:
        console.clear()
        display_banner()
        
        # Create a table that works well on both Windows and Mac
        table = Table(
            show_header=True,
            header_style="bold magenta",
            border_style="cyan",
            box=ASCII_BOX if sys.platform == "win32" else None
        )
        table.add_column("#", style="cyan", justify="right")
        table.add_column("Command", style="green")
        table.add_column("Description", style="white")
        
        for num, cmd, desc in menu_items:
            table.add_row(num, cmd, desc)
        
        console.print(table)
        
        try:
            choice = Prompt.ask("\n[yellow]Enter a number to select a command[/yellow]", default="1")
            if not choice.isdigit() or int(choice) < 1 or int(choice) > len(menu_items):
                console.print(f"[red]Invalid choice. Please enter a number between 1 and {len(menu_items)}[/red]")
                continue
            
            index = int(choice) - 1
            _, command_name, _ = menu_items[index]
            
            if command_name == "exit":
                console.print("[yellow]Goodbye![/yellow]")
                break
            elif command_name == "chat":
                # Handle chat command with model selection
                models_list = get_available_models()
                if not models_list:
                    console.print("[yellow]No models available. Is Ollama running?[/yellow]")
                    continue
                
                console.print("\n[cyan]Available models:[/cyan]")
                for i, model in enumerate(models_list, 1):
                    console.print(f"[green]{i}[/green]. {model}")
                
                model_choice = Prompt.ask("\nSelect model number", default="1")
                if not model_choice.isdigit() or int(model_choice) < 1 or int(model_choice) > len(models_list):
                    console.print("[red]Invalid model selection[/red]")
                    continue
                
                selected_model = models_list[int(model_choice) - 1]
                # Pass the model name directly as a string
                interactive_chat(model=selected_model, system_prompt=None, context_files=None)
            elif command_name == "pull":
                # Handle pull command
                model_name = Prompt.ask("\nEnter model name to pull")
                pull(model=model_name)
            else:
                # Execute the command
                command_func = commands.get(command_name)
                if command_func:
                    command_func()
            
            # Pause before showing menu again (except for settings which has its own pause)
            if command_name != "settings":
                Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")
        except KeyboardInterrupt:
            console.print("\n[yellow]Operation cancelled. Returning to menu...[/yellow]")
            continue
        except Exception as e:
            console.print(f"\n[red]Error: {str(e)}[/red]")
            Prompt.ask("\n[yellow]Press Enter to continue[/yellow]")

def perform_web_search(query: str) -> tuple[list, str]:
    """Perform a web search and return results with content"""
    try:
        print("Starting search...")  # Debug print
        with DDGS() as ddgs:
            # Get search results
            search_results = []
            raw_results = list(ddgs.text(query, max_results=10))
            print(f"Raw results: {raw_results}")  # Debug print
            
            for r in raw_results:
                print(f"Processing result: {r}")  # Debug print
                search_results.append({
                    "title": r.get("title", "No title"),
                    "link": r.get("link", ""),
                    "snippet": r.get("body", "No description available")
                })
            
            if not search_results:
                return [], ""
            
            # Prepare context for LLM
            context = "Here are the search results:\n\n"
            for i, result in enumerate(search_results, 1):
                context += f"{i}. {result['title']}\n"
                context += f"   URL: {result['link']}\n"
                context += f"   Summary: {result['snippet']}\n\n"
            
            return search_results, context
    except Exception as e:
        print(f"Search error: {str(e)}")  # Debug print
        console.print(f"[red]Error performing web search: {str(e)}[/red]")
        return [], ""

def fetch_webpage_content(url: str) -> str:
    """Fetch and extract text content from a webpage"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()  # Raise an error for bad status codes
        
        if response.status_code == 200:
            # Try to detect the encoding
            if 'charset' in response.headers.get('content-type', '').lower():
                response.encoding = response.apparent_encoding
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove unwanted elements
            for element in soup(['script', 'style', 'header', 'footer', 'nav']):
                element.decompose()
            
            # Convert HTML to markdown-like text
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.ignore_images = True
            h.ignore_tables = True
            h.body_width = 0  # Don't wrap text
            text = h.handle(str(soup))
            
            # Clean up the text
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            text = ' '.join(lines)
            
            # Remove multiple spaces and normalize whitespace
            text = ' '.join(text.split())
            
            # Limit text length
            return text[:2000]  # Limit to first 2000 chars to avoid overloading the LLM
    except requests.RequestException as e:
        console.print(f"[dim red]Network error fetching {url}: {str(e)}[/dim red]")
    except Exception as e:
        console.print(f"[dim red]Error processing {url}: {str(e)}[/dim red]")
    return ""

def process_enhanced_search(query: str, model: str) -> str:
    """Process an enhanced search query with web results and return formatted analysis"""
    try:
        # Strip the "search:" prefix if present
        search_query = query.replace("search:", "", 1).strip()
        
        with console.status("[cyan]Searching the web...[/cyan]"):
            results, context = perform_web_search(search_query)
        
        if not results:
            return "No search results found for your query. Please try a different search term."
        
        # Prepare system prompt for analysis
        system_prompt = """You are a helpful AI assistant that analyzes search results and provides comprehensive summaries. 
        For the given query, provide a well-structured analysis including:
        1. A direct answer or summary addressing the main query
        2. Key findings and insights from the search results
        3. Supporting evidence from multiple sources where available
        4. Any important caveats, limitations, or conflicting information
        5. Relevant quotes or specific details that support the analysis
        
        Format your response in clear sections with markdown headings and bullet points for readability.
        Base your response solely on the provided search results."""
        
        # Prepare the prompt for the LLM
        prompt = f"""
        Query: {search_query}

        Please analyze these search results and provide a comprehensive summary following the structure outlined in the system prompt.

        Search Results:
        {context}
        """
        
        # Send to Ollama using the /chat endpoint for consistency with Jira integration
        with console.status("[cyan]Analyzing search results...[/cyan]"):
            response = requests.post(
                f"{OLLAMA_API}/chat",
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": "You are a helpful assistant that analyzes web search results."},
                        {"role": "user", "content": prompt}
                    ],
                    "stream": False
                }
            )
            
            if response.status_code != 200:
                return f"Error: Failed to analyze search results. Status code: {response.status_code}"
            
            result = response.json().get("response", "")
            if not result:
                return "Error: No analysis was generated. Please try again."
            
            return result
            
    except requests.RequestException as e:
        return f"Network error occurred while processing search: {str(e)}"
    except Exception as e:
        return f"An error occurred while processing search: {str(e)}"

@app.command()
def analyze(
    image: str = typer.Argument(..., help="Path to the image file to analyze"),
    prompt: Optional[str] = typer.Option(None, help="Custom prompt for image analysis")
):
    """Analyze an image using the vision model"""
    try:
        result = analyze_image(image, prompt)
        console.print("\n[bold cyan]Image Analysis:[/bold cyan]")
        console.print(Panel(result, border_style="cyan"))
    except Exception as e:
        console.print(f"[red]Error: {str(e)}[/red]")

@app.command()
def confluence(
    command: str = typer.Argument(None, help="Natural language command for Confluence operations"),
    mode: bool = typer.Option(False, "--mode", "-m", help="Enter interactive Confluence mode")
):
    """Execute Confluence operations using natural language or enter Confluence mode"""
    try:
        if not CONFLUENCE_MCP_AVAILABLE:
            console.print("[red]Confluence MCP Protocol integration not available.[/red]")
            console.print("[yellow]Please ensure the required dependencies are installed.[/yellow]")
            return
        
        # Check if mode flag is set
        if mode:
            confluence_mode()
            return
        
        # Check if Confluence is configured
        is_configured, message = check_confluence_configuration()
        if not is_configured:
            console.print(f"[red]Confluence is not configured: {message}[/red]")
            
            # Offer to configure Confluence
            setup_now = Prompt.ask("Would you like to set up Confluence integration now?", choices=["y", "n"], default="y")
            if setup_now.lower() == "y":
                # Get Confluence configuration
                console.print("\n[cyan]Confluence Configuration[/cyan]")
                console.print("You'll need your Confluence Cloud URL and API token to proceed.")
                console.print("You can generate an API token at https://id.atlassian.com/manage-profile/security/api-tokens")
                
                confluence_url = Prompt.ask("Enter your Confluence Cloud URL (e.g., https://your-domain.atlassian.net)")
                confluence_email = Prompt.ask("Enter your Confluence email address")
                confluence_token = Prompt.ask("Enter your Confluence API token", password=True)
                
                # Save configuration
                if save_confluence_config(confluence_url, confluence_token, confluence_email):
                    console.print("[green]Confluence configuration saved successfully![/green]")
                    is_configured = True
                else:
                    console.print("[red]Failed to save Confluence configuration.[/red]")
                    return
            else:
                return
        
        # Get the command if not provided
        if not command:
            command = Prompt.ask("[cyan]Enter your natural language Confluence command[/cyan]")
        
        # Get config for model
        config = load_config()
        current_model = config.get("default_model", "llama3")
        
        # Show a spinner while processing
        with Progress(
            SpinnerColumn(),
            TextColumn("[cyan]Processing Confluence command...[/cyan]"),
            transient=True
        ) as progress:
            progress.add_task("Processing", total=None)
            
            # Handle natural language Confluence command
            result = handle_confluence_nl_command(command, model=current_model)
        
        # Display the result
        display_confluence_result(result)
        
    except Exception as e:
        console.print(f"[red]Error: {str(e)}[/red]")

@app.command()
def jira(
    command: str = typer.Argument(None, help="Natural language command for Jira operations"),
    mode: bool = typer.Option(False, "--mode", "-m", help="Enter interactive Jira mode")
):
    """Execute Jira operations using natural language or enter Jira mode"""
    try:
        if not JIRA_MCP_AVAILABLE:
            console.print("[red]Jira MCP Protocol integration not available.[/red]")
            console.print("[yellow]Please ensure the required dependencies are installed.[/yellow]")
            return
        
        # Check if mode flag is set
        if mode:
            jira_mode()
            return
        
        # Check if Jira is configured
        is_configured, message = check_jira_configuration()
        if not is_configured:
            console.print(f"[red]Jira is not configured: {message}[/red]")
            
            # Offer to configure Jira
            setup_now = Prompt.ask("Would you like to set up Jira integration now?", choices=["y", "n"], default="y")
            if setup_now.lower() == "y":
                # Get Jira configuration
                console.print("\n[cyan]Jira Configuration[/cyan]")
                console.print("You'll need your Jira URL and API token to proceed.")
                console.print("You can generate an API token at https://id.atlassian.com/manage-profile/security/api-tokens")
                
                jira_url = Prompt.ask("Enter your Jira URL (e.g., https://your-domain.atlassian.net)")
                jira_email = Prompt.ask("Enter your Jira email address")
                jira_token = Prompt.ask("Enter your Jira API token", password=True)
                
                # Save configuration
                if save_jira_config(jira_url, jira_token, jira_email):
                    console.print("[green]Jira configuration saved successfully![/green]")
                    is_configured = True
                else:
                    console.print("[red]Failed to save Jira configuration.[/red]")
                    return
            else:
                return
        
        # Get the command if not provided
        if not command:
            command = Prompt.ask("[cyan]Enter your natural language Jira command[/cyan]")
        
        # Get config for model
        config = load_config()
        current_model = config.get("default_model", "llama3")
        
        # Show a spinner while processing
        with Progress(
            SpinnerColumn(),
            TextColumn("[cyan]Processing Jira command...[/cyan]"),
            transient=True
        ) as progress:
            progress.add_task("Processing", total=None)
            
            # Handle natural language Jira command
            result = handle_jira_nl_command(command, model=current_model)
        
        # Display the result
        display_jira_result(result)
        
    except Exception as e:
        console.print(f"[red]Error: {str(e)}[/red]")

def analyze_from_menu():
    """Handle image analysis from the menu"""
    try:
        # First check if vision model is available
        config = load_config()
        vision_model = config.get("default_vision_model", "llama3.2-vision:latest")
        models_list = get_available_models()
        
        if vision_model not in models_list:
            console.print(f"[yellow]Warning: Vision model {vision_model} not found.[/yellow]")
            pull_model = Prompt.ask(f"Would you like to pull {vision_model} now? (y/n)", choices=["y", "n", "yes", "no"], default="y")
            if pull_model.lower() in ['y', 'yes']:
                pull(vision_model)
            else:
                console.print("[red]Cannot analyze images without a vision model.[/red]")
                return
        
        image_path = Prompt.ask("[cyan]Enter the path to the image file[/cyan]")
        image_path = os.path.expanduser(image_path)
        
        if not os.path.exists(image_path):
            console.print("[red]Error: File not found[/red]")
            return
        
        use_custom_prompt = Prompt.ask("[cyan]Would you like to use a custom prompt? (y/n)[/cyan]", choices=["y", "n", "yes", "no"], default="n")
        prompt = None
        if use_custom_prompt.lower() in ['y', 'yes']:
            prompt = Prompt.ask("[cyan]Enter your custom prompt[/cyan]")
        
        result = analyze_image(image_path, prompt)
        console.print("\n[bold cyan]Image Analysis:[/bold cyan]")
        console.print(Panel(result, border_style="cyan"))
    except Exception as e:
        console.print(f"[red]Error: {str(e)}[/red]")

def encode_image(image_path: str) -> tuple[str, str]:
    """
    Encode an image file to base64 and determine its MIME type.
    Returns a tuple of (base64_string, mime_type)
    """
    mime_type = mimetypes.guess_type(image_path)[0]
    if not mime_type or not mime_type.startswith('image/'):
        raise ValueError(f"Unsupported file type: {image_path}")
    
    try:
        with Image.open(image_path) as img:
            # Convert to RGB if necessary
            if img.mode in ('RGBA', 'LA'):
                img = img.convert('RGB')
            
            # Save to BytesIO in appropriate format
            img_byte_arr = BytesIO()
            img.save(img_byte_arr, format=img.format or 'JPEG')
            img_byte_arr.seek(0)
            
            return base64.b64encode(img_byte_arr.getvalue()).decode('utf-8'), mime_type
    except Exception as e:
        raise Exception(f"Error processing image {image_path}: {str(e)}")

def analyze_image(image_path: str, prompt: Optional[str] = None) -> str:
    """
    Analyze an image using the vision model.
    Args:
        image_path: Path to the image file
        prompt: Optional prompt to guide the analysis
    Returns:
        Analysis result from the model
    """
    try:
        config = load_config()
        model = config.get("default_vision_model", "llama3.2-vision:latest")
        
        # First verify the model exists
        models_list = get_available_models()
        if model not in models_list:
            raise Exception(f"Vision model {model} not found. Please pull it first using: ollama pull {model}")
        
        # Encode image
        base64_image, mime_type = encode_image(image_path)
        
        # Format according to Ollama vision model spec
        default_prompt = "Please analyze this image and describe what you see in detail."
        user_prompt = prompt or default_prompt
        
        # Send to Ollama using the /chat endpoint for consistency with Jira integration
        with console.status(f"[cyan]Analyzing image with {model}...[/cyan]"):
            response = requests.post(
                f"{OLLAMA_API}/chat",
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": "You are a helpful assistant that analyzes images."},
                        {"role": "user", "content": user_prompt, "images": [base64_image]}
                    ],
                    "stream": False
                }
            )
            
            if response.status_code != 200:
                raise Exception(f"Error from Ollama API: {response.text}")
            
            # Try to parse the response as JSON
            try:
                result_json = response.json()
                if "response" in result_json:
                    return result_json["response"]
                else:
                    # Try alternative response format
                    if "message" in result_json and "content" in result_json["message"]:
                        return result_json["message"]["content"]
                    else:
                        raise Exception("Unexpected response format from model")
            except json.JSONDecodeError:
                # If JSON parsing fails, return the raw text
                result = response.text.strip()
                if result:
                    return result
                else:
                    raise Exception("No analysis received from model")
            
    except Exception as e:
        raise Exception(f"Error analyzing image: {str(e)}")


# display_confluence_result function is imported from ollama_shell_confluence_mcp

def confluence_mode():
    """
    Enter a Confluence mode where users can interact with Confluence Cloud using natural language.
    This provides a simple interface to access Confluence content within Ollama Shell.
    """
    if not CONFLUENCE_MCP_AVAILABLE:
        console.print("[red]Confluence MCP integration is not available.[/red]")
        console.print("[yellow]Please ensure the required dependencies are installed.[/yellow]")
        return
    
    # Check if Confluence is configured
    is_configured, message = check_confluence_configuration()
    if not is_configured:
        console.print(f"[red]{message}[/red]")
        
        # Offer to configure Confluence
        setup_now = Prompt.ask("Would you like to set up Confluence integration now?", choices=["y", "n"], default="y")
        if setup_now.lower() == "y":
            # Get Confluence configuration
            console.print("\n[cyan]Confluence Configuration[/cyan]")
            console.print("You'll need your Confluence Cloud URL and API token to proceed.")
            console.print("You can generate an API token at https://id.atlassian.com/manage-profile/security/api-tokens")
            
            confluence_url = Prompt.ask("Enter your Confluence Cloud URL (e.g., https://your-domain.atlassian.net)")
            confluence_email = Prompt.ask("Enter your Confluence email address")
            confluence_token = Prompt.ask("Enter your Confluence API token", password=True)
            
            # Save configuration
            if save_confluence_config(confluence_url, confluence_token, confluence_email):
                console.print("[green]Confluence configuration saved successfully![/green]")
                is_configured = True
            else:
                console.print("[red]Failed to save Confluence configuration.[/red]")
                return
        else:
            return
    
    # Main Confluence interaction loop
    console.clear()
    display_banner()
    
    # Get default model
    config = load_config()
    default_model = config.get("default_model", "llama3")
    
    # Create a table with available commands
    table = Table(title="Confluence Commands")
    table.add_column("Command", style="cyan")
    table.add_column("Description", style="white")
    
    table.add_row("/help", "Show this help message")
    table.add_row("/model <model_name>", "Change the model used for Confluence queries")
    table.add_row("/exit", "Exit Confluence mode")
    table.add_row("Any other text", "Will be interpreted as a natural language query to Confluence")
    
    console.print(table)
    console.print(f"\n[green]Using model:[/green] {default_model}")
    console.print("[cyan]Enter your Confluence queries below. Type /exit to return to the main menu.[/cyan]")
    
    # Set up prompt session with key bindings
    kb = KeyBindings()
    session = PromptSession(key_bindings=kb)
    
    current_model = default_model
    
    while True:
        try:
            # Get user input
            user_input = session.prompt("\n[Confluence] > ")
            
            # Handle commands
            if user_input.lower() == "/exit":
                break
            elif user_input.lower() == "/help":
                console.print(table)
                continue
            elif user_input.lower().startswith("/model "):
                # Change model
                model_name = user_input[7:].strip()
                is_valid, error_message = validate_model(model_name)
                if is_valid:
                    current_model = model_name
                    console.print(f"[green]Model changed to {current_model}[/green]")
                else:
                    console.print(f"[red]Error: {error_message}[/red]")
                continue
            
            # Process natural language query to Confluence
            if user_input.strip():
                with Progress(
                    SpinnerColumn(),
                    TextColumn("[bold blue]Processing Confluence query..."),
                    transient=True
                ) as progress:
                    progress.add_task("Processing", total=None)
                    result = handle_confluence_nl_command(user_input, model=current_model)
                
                # Display result
                display_confluence_result(result)
        
        except KeyboardInterrupt:
            console.print("\n[yellow]Operation cancelled.[/yellow]")
            continue
        except Exception as e:
            console.print(f"\n[red]Error: {str(e)}[/red]")

def jira_mode():
    """
    Enter a Jira mode where users can interact with Jira using natural language.
    This provides a simple interface to access Jira issues within Ollama Shell.
    """
    if not JIRA_MCP_AVAILABLE:
        console.print("[red]Jira MCP integration is not available.[/red]")
        console.print("[yellow]Please ensure the required dependencies are installed.[/yellow]")
        return
    
    # Check if Jira is configured
    is_configured, message = check_jira_configuration()
    if not is_configured:
        console.print(f"[red]{message}[/red]")
        
        # Offer to configure Jira
        setup_now = Prompt.ask("Would you like to set up Jira integration now?", choices=["y", "n"], default="y")
        if setup_now.lower() == "y":
            # Get Jira configuration
            console.print("\n[cyan]Jira Configuration[/cyan]")
            console.print("You'll need your Jira URL and API token to proceed.")
            console.print("You can generate an API token at https://id.atlassian.com/manage-profile/security/api-tokens")
            
            jira_url = Prompt.ask("Enter your Jira URL (e.g., https://your-domain.atlassian.net)")
            jira_email = Prompt.ask("Enter your Jira email address")
            jira_token = Prompt.ask("Enter your Jira API token", password=True)
            
            # Save configuration
            if save_jira_config(jira_url, jira_token, jira_email):
                console.print("[green]Jira configuration saved successfully![/green]")
                is_configured = True
            else:
                console.print("[red]Failed to save Jira configuration.[/red]")
                return
        else:
            return
    
    # Main Jira interaction loop
    console.clear()
    display_banner()
    
    # Get default model
    config = load_config()
    default_model = config.get("default_model", "llama3")
    
    # Create a table with available commands
    table = Table(title="Jira Commands")
    table.add_column("Command", style="cyan")
    table.add_column("Description", style="white")
    
    table.add_row("/help", "Show this help message")
    table.add_row("/model <model_name>", "Change the model used for Jira queries")
    table.add_row("/exit", "Exit Jira mode")
    table.add_row("/search <query>", "Search for issues using JQL or natural language")
    table.add_row("/get <issue-key>", "Get details of a specific issue")
    table.add_row("/comment <issue-key> <comment>", "Add a comment to an issue")
    table.add_row("/update <issue-key> <field> <value>", "Update an issue field")
    table.add_row("Any other text", "Will be interpreted as a natural language query to Jira")
    
    console.print(table)
    console.print(f"\n[green]Using model:[/green] {default_model}")
    console.print("[cyan]Enter your Jira queries below. Type /exit to return to the main menu.[/cyan]")
    
    # Set up prompt session with key bindings
    kb = KeyBindings()
    session = PromptSession(key_bindings=kb)
    
    current_model = default_model
    
    while True:
        try:
            # Get user input
            user_input = session.prompt("\n[Jira] > ")
            
            # Handle commands
            if user_input.lower() == "/exit":
                break
            elif user_input.lower() == "/help":
                console.print(table)
                continue
            elif user_input.lower().startswith("/model "):
                # Change model
                model_name = user_input[7:].strip()
                is_valid, error_message = validate_model(model_name)
                if is_valid:
                    current_model = model_name
                    console.print(f"[green]Model changed to {current_model}[/green]")
                else:
                    console.print(f"[red]Error: {error_message}[/red]")
                continue
            
            # Process natural language query to Jira
            if user_input.strip():
                with Progress(
                    SpinnerColumn(),
                    TextColumn("[bold blue]Processing Jira query..."),
                    transient=True
                ) as progress:
                    progress.add_task("Processing", total=None)
                    result = handle_jira_nl_command(user_input, model=current_model)
                
                # Display result
                display_jira_result(result)
        
        except KeyboardInterrupt:
            console.print("\n[yellow]Operation cancelled.[/yellow]")
            continue
        except Exception as e:
            console.print(f"\n[red]Error: {str(e)}[/red]")

def terminal_mode():
    """
    Enter a terminal mode where users can execute system commands.
    This provides a simple shell interface within Ollama Shell.
    """
    import subprocess
    import shlex
    import threading
    import select
    import time
    import signal
    
    console.print("\n[bold cyan]Terminal Mode[/bold cyan]")
    console.print("[yellow]Type commands to execute them. Type 'exit' to return to the main menu.[/yellow]")
    console.print("[yellow]Press Ctrl+C to interrupt a running command.[/yellow]")
    console.print("[yellow]Current directory: [green]{0}[/green][/yellow]".format(os.getcwd()))
    
    # Track the current process for handling interrupts
    current_process = None
    
    def interrupt_handler(sig, frame):
        if current_process:
            try:
                # On Unix systems, send SIGINT to the process group
                if os.name != 'nt':  # Not Windows
                    os.killpg(os.getpgid(current_process.pid), signal.SIGINT)
                else:
                    current_process.terminate()
                console.print("\n[yellow]Command interrupted[/yellow]")
            except Exception as e:
                console.print(f"\n[red]Error interrupting process: {str(e)}[/red]")
    
    # Set up signal handler
    original_handler = signal.getsignal(signal.SIGINT)
    signal.signal(signal.SIGINT, interrupt_handler)
    
    try:
        while True:
            try:
                # Get command from user
                cmd = Prompt.ask("\n[bold green]$[/bold green]")
                
                # Exit terminal mode
                if cmd.lower() in ['exit', 'quit', 'q']:
                    break
                    
                # Handle cd command specially since it affects the Python process
                if cmd.startswith('cd '):
                    try:
                        # Extract the directory path
                        directory = cmd[3:].strip()
                        # Expand user directory if needed (e.g., ~/Documents)
                        directory = os.path.expanduser(directory)
                        # Change directory
                        os.chdir(directory)
                        console.print("[yellow]Changed directory to: [green]{0}[/green][/yellow]".format(os.getcwd()))
                    except Exception as e:
                        console.print(f"[red]Error: {str(e)}[/red]")
                    continue
                    
                # Handle clear command
                if cmd.lower() in ['clear', 'cls']:
                    console.clear()
                    console.print("[bold cyan]Terminal Mode[/bold cyan]")
                    console.print("[yellow]Type commands to execute them. Type 'exit' to return to the main menu.[/yellow]")
                    console.print("[yellow]Press Ctrl+C to interrupt a running command.[/yellow]")
                    console.print("[yellow]Current directory: [green]{0}[/green][/yellow]".format(os.getcwd()))
                    continue
                
                # Execute the command with real-time output
                try:
                    # Use shell=False for security and proper argument splitting
                    # Create a new process group on Unix systems
                    if os.name != 'nt':  # Not Windows
                        current_process = subprocess.Popen(
                            shlex.split(cmd),
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            text=True,
                            shell=False,
                            bufsize=1,  # Line buffered
                            preexec_fn=os.setsid  # Create new process group
                        )
                    else:  # Windows
                        current_process = subprocess.Popen(
                            shlex.split(cmd),
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            text=True,
                            shell=False,
                            bufsize=1  # Line buffered
                        )
                    
                    # Function to read from a pipe and print output in real-time
                    def read_pipe(pipe, is_error=False):
                        while True:
                            line = pipe.readline()
                            if not line:
                                break
                            if is_error:
                                console.print(f"[red]{line}[/red]", end="")
                            else:
                                console.print(line, end="")
                    
                    # Create threads to read stdout and stderr
                    stdout_thread = threading.Thread(target=read_pipe, args=(current_process.stdout, False))
                    stderr_thread = threading.Thread(target=read_pipe, args=(current_process.stderr, True))
                    
                    # Start threads
                    stdout_thread.daemon = True
                    stderr_thread.daemon = True
                    stdout_thread.start()
                    stderr_thread.start()
                    
                    # Wait for the process to complete
                    return_code = current_process.wait()
                    
                    # Wait for output threads to finish
                    stdout_thread.join()
                    stderr_thread.join()
                    
                    # Print return code if non-zero
                    if return_code != 0:
                        console.print(f"[yellow]Command exited with code: {return_code}[/yellow]")
                    
                    # Reset current process
                    current_process = None
                        
                except Exception as e:
                    console.print(f"[red]Error executing command: {str(e)}[/red]")
                    current_process = None
                    
            except KeyboardInterrupt:
                # This should be handled by our signal handler
                pass
            except EOFError:
                break
    finally:
        # Restore original signal handler
        signal.signal(signal.SIGINT, original_handler)

if __name__ == "__main__":
    if len(sys.argv) == 1:
        display_menu()
    else:
        app()
